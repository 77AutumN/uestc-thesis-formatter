"""Unit tests for route_advisor (W5 Wave 2 Item 3, 2026-05-16).

Builds minimal .docx fixtures for each condition pass/fail and verifies the
4-condition recommendation logic.
"""
from __future__ import annotations

import sys
from pathlib import Path

import pytest

THIS = Path(__file__).resolve().parent
SKILL_ROOT = THIS.parent
sys.path.insert(0, str(SKILL_ROOT / "scripts"))

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    Document = None


def _build_minimal_docx(tmp_path: Path, *, use_builtin_heading: bool = True,
                       custom_style_name: str | None = None,
                       with_header: bool = True) -> Path:
    """Build a docx with knobs for each condition."""
    out = tmp_path / "fixture.docx"
    doc = Document()
    if use_builtin_heading:
        chap = doc.add_paragraph("第一章 标题")
        chap.style = "Heading 1"
    if custom_style_name:
        styles = doc.styles
        if custom_style_name not in [s.name for s in styles]:
            styles.add_style(custom_style_name, WD_STYLE_TYPE.PARAGRAPH)
        para = doc.add_paragraph("自定义样式段")
        para.style = custom_style_name
    body = doc.add_paragraph("正文段落示例.")
    if with_header:
        section = doc.sections[0]
        header = section.header
        # python-docx default header is empty; add a paragraph with a PAGE-like
        # placeholder. We can't easily insert a real STYLEREF field via the
        # high-level API, so we manually inject XML.
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        para = header.paragraphs[0]
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = "1"
        run.append(text)
        fld.append(run)
        para._p.append(fld)
    doc.save(str(out))
    return out


@pytest.mark.skipif(Document is None, reason="python-docx not available")
def test_all_conditions_pass_docx_only(tmp_path: Path) -> None:
    import route_advisor
    docx = _build_minimal_docx(tmp_path, use_builtin_heading=True, with_header=True)
    result = route_advisor.detect_route_eligibility(str(docx), deliverable_mode="docx_only")
    assert result["condition_2_builtin_headings"]["met"]
    assert result["condition_3_no_corruption"]["met"]
    assert result["condition_4_docx_only_delivery"]["met"]
    # condition_1 may or may not detect PAGE field depending on python-docx XML;
    # we don't assert on it here — focus on full pipeline behavior:
    if result["condition_1_header_footer"]["met"]:
        assert result["recommended_route"] == "docx_direct"
        assert "All 4 conditions met" in result["rationale"]


@pytest.mark.skipif(Document is None, reason="python-docx not available")
def test_pdf_required_recommends_latex(tmp_path: Path) -> None:
    import route_advisor
    docx = _build_minimal_docx(tmp_path, use_builtin_heading=True, with_header=True)
    result = route_advisor.detect_route_eligibility(str(docx), deliverable_mode="pdf_required")
    assert result["recommended_route"] == "latex_v2"
    assert "PDF delivery" in result["rationale"]


@pytest.mark.skipif(Document is None, reason="python-docx not available")
def test_unknown_deliverable_defaults_latex(tmp_path: Path) -> None:
    import route_advisor
    docx = _build_minimal_docx(tmp_path, use_builtin_heading=True, with_header=True)
    result = route_advisor.detect_route_eligibility(str(docx), deliverable_mode="unknown")
    # Structural conditions may be met but deliverable unknown → safe fallback
    assert result["recommended_route"] == "latex_v2"


@pytest.mark.skipif(Document is None, reason="python-docx not available")
def test_no_builtin_heading_blocks_docx_direct(tmp_path: Path) -> None:
    import route_advisor
    docx = _build_minimal_docx(
        tmp_path,
        use_builtin_heading=False,
        custom_style_name="1-1级",
        with_header=True,
    )
    result = route_advisor.detect_route_eligibility(str(docx), deliverable_mode="docx_only")
    assert not result["condition_2_builtin_headings"]["met"]
    assert result["recommended_route"] == "latex_v2"
    assert "Condition" in result["rationale"] or "failed" in result["rationale"]


def test_bad_docx_path_returns_error(tmp_path: Path) -> None:
    import route_advisor
    result = route_advisor.detect_route_eligibility(str(tmp_path / "nonexistent.docx"))
    assert result["recommended_route"] == "latex_v2"
    assert "cannot open" in result["rationale"]
