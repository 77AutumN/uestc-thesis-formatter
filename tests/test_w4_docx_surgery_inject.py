"""tests/test_w4_docx_surgery_inject.py — W4-C inject + delete_orphan 联合测试.

覆盖:
  - _detector_inject_heading_before: anchor 上移到章引言段
  - _detector_orphan_title_paragraph: 识别"裸章名段"
  - _apply_inject_heading_before: 真实 inject 段
  - _apply_delete_orphan_title_paragraph: 真实 delete 段

case19 round 2/2c 凝结的 4 项 surgery 在 W4-C 一次完成.
"""
from __future__ import annotations
import os
import shutil
import sys
import tempfile

import pytest

THIS = os.path.dirname(os.path.abspath(__file__))
SCRIPTS = os.path.normpath(os.path.join(THIS, "..", "scripts"))
if SCRIPTS not in sys.path:
    sys.path.insert(0, SCRIPTS)

import docx_surgery as ds  # noqa: E402
import source_manifest as sm  # noqa: E402

FX_INJECT_ORPHAN = os.path.join(THIS, "fixtures", "w4_inject_orphan.docx")


def _need(p):
    if not os.path.isfile(p):
        pytest.skip(f"fixture missing: {p} (run build_w4_inject_orphan.py)")


# ============================================================
# Detector tests
# ============================================================

def test_inject_detector_early_exit_when_prefixes_present():
    """fixture body 已含"第N章 X" prefix 段, source_manifest 自动判 H1 accepted →
    detector 应早退返回 []. 真实 inject 触发场景是 case14 (custom style + 缺 prefix),
    需要更智能的章号推断, 不在 W4-C 主体范围."""
    _need(FX_INJECT_ORPHAN)
    manifest = sm.build_probe_manifest(FX_INJECT_ORPHAN)
    h1_accepted = [h for h in manifest["headings"]
                   if h.get("level") == 1 and h.get("status") == "accepted"]
    assert len(h1_accepted) > 0, "fixture should produce H1 from para_regex"
    ops = ds._detector_inject_heading_before(manifest)
    assert ops == [], "detector should early-exit when body already has H1"


def test_inject_detector_anchor_preface_logic_unit():
    """anchor 上移逻辑单元测试: 用 fake manifest 模拟 0 H1 + body 第N章 + 引言段."""
    fake_manifest = {
        "headings": [],  # 0 H1 → 不早退
        "paragraphs": [
            {"id": "p001", "text": "前一章末尾段, 是上一章的结论或最后内容, 长度足够触发引言段判定 (>30字符).",
             "zone_guess": "body", "style_name": "Normal"},
            {"id": "p_blank", "text": "", "zone_guess": "body", "style_name": "Normal"},  # chapter 边界空段
            {"id": "p002", "text": "针对当前测试方法工具组合存在的短板, 结合应用场景的实际情况, 本文提出以下优化对策建议, 推动测试方法体系的进一步完善.",
             "zone_guess": "body", "style_name": "Normal"},  # 章引言段
            {"id": "p003", "text": "第六章 测试方法体系优化路径",
             "zone_guess": "body", "style_name": "Normal"},  # body 第N章 段
            {"id": "p004", "text": "6.1 优化方法内部结构", "zone_guess": "body", "style_name": "Normal"},
        ],
    }
    ops = ds._detector_inject_heading_before(fake_manifest)
    assert len(ops) == 1
    op = ops[0]
    # anchor 应上移到引言段 p002, 不是 p003 (第N章段 itself)
    assert op["params"]["anchor_paragraph_id"] == "p002"
    assert "针对当前测试方法" in op["params"]["anchor_text_match"]
    assert op["visible_text_change"] is True
    assert op["params"]["title"] == "第六章 测试方法体系优化路径"


def test_orphan_detector_finds_naked_chapter_name():
    """裸章名段 '测试方法体系优化路径' 应被识别为 delete 候选."""
    _need(FX_INJECT_ORPHAN)
    manifest = sm.build_probe_manifest(FX_INJECT_ORPHAN)
    ops = ds._detector_orphan_title_paragraph(manifest)
    targets = [op["params"]["target_text_match"] for op in ops]
    assert any("测试方法体系优化路径" == t for t in targets), \
        f"expected naked chapter name in targets, got: {targets}"


def test_orphan_op_marks_visible_text_change_true():
    """delete_orphan 删可见文字, 必须 visible_text_change=True."""
    _need(FX_INJECT_ORPHAN)
    manifest = sm.build_probe_manifest(FX_INJECT_ORPHAN)
    ops = ds._detector_orphan_title_paragraph(manifest)
    if not ops:
        pytest.skip("no orphan ops detected, schema check skipped")
    assert all(op["visible_text_change"] is True for op in ops)


# ============================================================
# Apply tests (real docx mutation)
# ============================================================

def test_apply_inject_heading_before_inserts_paragraph():
    """apply 在 anchor 段之前插入 Heading 1 段."""
    _need(FX_INJECT_ORPHAN)
    with tempfile.TemporaryDirectory() as td:
        tmp = os.path.join(td, "case.docx")
        shutil.copy(FX_INJECT_ORPHAN, tmp)
        params = {
            "anchor_paragraph_id": "?",
            "anchor_text_match": "6.1 优化方法内部结构",
            "title": "TEST 第六章 X",
            "level": 1,
        }
        result = ds._apply_inject_heading_before(tmp, params)
        assert result["paragraphs_added"] == 1

        from docx import Document
        d = Document(tmp)
        texts = [p.text for p in d.paragraphs]
        # inject 段应当出现在 anchor 段之前
        idx_inject = next(i for i, t in enumerate(texts) if t == "TEST 第六章 X")
        idx_anchor = next(i for i, t in enumerate(texts) if t == "6.1 优化方法内部结构")
        assert idx_inject == idx_anchor - 1, f"inject @ {idx_inject}, anchor @ {idx_anchor}"


def test_apply_delete_orphan_removes_paragraph():
    """apply 删除指定段."""
    _need(FX_INJECT_ORPHAN)
    with tempfile.TemporaryDirectory() as td:
        tmp = os.path.join(td, "case.docx")
        shutil.copy(FX_INJECT_ORPHAN, tmp)
        params = {
            "target_paragraph_id": "?",
            "target_text_match": "测试方法体系优化路径",
        }
        result = ds._apply_delete_orphan_title_paragraph(tmp, params)
        assert result["paragraphs_deleted"] == 1

        from docx import Document
        d = Document(tmp)
        texts = [p.text.strip() for p in d.paragraphs]
        # 裸段已删, "第六章 X" 全 prefix 段应仍在
        assert "测试方法体系优化路径" not in texts
        assert any("第六章" in t and "测试方法体系优化路径" in t for t in texts)


def test_apply_inject_missing_anchor_raises():
    """anchor text 找不到 → RuntimeError."""
    _need(FX_INJECT_ORPHAN)
    with tempfile.TemporaryDirectory() as td:
        tmp = os.path.join(td, "case.docx")
        shutil.copy(FX_INJECT_ORPHAN, tmp)
        params = {
            "anchor_paragraph_id": "?",
            "anchor_text_match": "NONEXISTENT ANCHOR TEXT",
            "title": "X",
            "level": 1,
        }
        with pytest.raises(RuntimeError, match="anchor not found"):
            ds._apply_inject_heading_before(tmp, params)


def test_apply_delete_missing_target_raises():
    """target text 找不到 → RuntimeError."""
    _need(FX_INJECT_ORPHAN)
    with tempfile.TemporaryDirectory() as td:
        tmp = os.path.join(td, "case.docx")
        shutil.copy(FX_INJECT_ORPHAN, tmp)
        params = {"target_paragraph_id": "?", "target_text_match": "NONEXISTENT"}
        with pytest.raises(RuntimeError, match="orphan target not found"):
            ds._apply_delete_orphan_title_paragraph(tmp, params)
