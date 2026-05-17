"""W4-C inject_heading_before + delete_orphan_title_paragraph 联合 fixture.

模拟 case19 round 2/2b 的 docx 结构:
  - body 含 "第二章 X" / "第四章 Y" 标准 chapter heading 段
  - 但 "第六章 Z" 之前有客户手写"裸章名段" Z (无前缀, 与 inject 重复)
  - 还有一段章引言, anchor 应上移到引言段之前

期望 surgery 自动:
  1. inject "第六章 Z" 在章引言段之前
  2. delete "Z" 裸章名段
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent


def build_w4_inject_orphan():
    out = HERE / "w4_inject_orphan.docx"
    if out.exists():
        out.unlink()
    d = Document()
    # 第二章 (有 prefix) — 不需 surgery
    d.add_paragraph("第二章 理论基础与核心概念")
    d.add_paragraph("2.1 核心概念界定")
    d.add_paragraph("本章先界定核心概念, 然后梳理理论, 系统呈现研究进展. 内容长度足以触发引言段判定 (>30字符).")
    d.add_paragraph("2.1.1 子概念")
    d.add_paragraph("2.1.1 正文段, 长度足以让 sentence_like 启发命中.")

    # 第四章 (有 prefix) — 不需 surgery
    d.add_paragraph("第四章 测试方法特征与现存问题")
    d.add_paragraph("4.1 数据采集方法")

    # ===== 关键场景: 第六章 inject-orphan 模式 =====
    # 客户手写"裸章名段" (无 "第六章" prefix) — 需要 detector 标 delete
    d.add_paragraph("测试方法体系优化路径")
    # 章引言段 — anchor 应上移到这里
    d.add_paragraph("针对当前测试方法工具组合存在的短板, 结合应用场景的实际情况, 本文提出以下优化对策建议, 推动测试方法体系的进一步完善.")
    # "第六章 Z" 章标题段 (有 prefix) — detector 找到此, anchor 应上移到上面引言
    d.add_paragraph("第六章 测试方法体系优化路径")
    d.add_paragraph("6.1 优化方法内部结构")
    d.add_paragraph("6.1 正文段足够长触发启发判定阈值, 这是用于 zone_guess 的判断.")

    d.save(out)
    print(f"wrote {out} ({out.stat().st_size} bytes)")


if __name__ == "__main__":
    build_w4_inject_orphan()
