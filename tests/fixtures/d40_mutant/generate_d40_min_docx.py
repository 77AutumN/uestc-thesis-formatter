"""Generate a minimal D40 docx mutant fixture.

The document contains:
  - one Heading 1 chapter anchor
  - one whole-paragraph inline-math + numbered suffix trigger
  - one control paragraph with inline math but no trailing number

Usage:
    python generate_d40_min_docx.py [--output PATH]
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


DEFAULT_OUTPUT = Path(__file__).with_name("d40_min.docx")


def build_docx(output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    chapter = doc.add_paragraph("第一章 D40 试验")
    chapter.style = "Heading 1"

    trigger = doc.add_paragraph()
    trigger.add_run("$v = at$ (1-2)")

    control = doc.add_paragraph()
    control.add_run("对照段 ")
    control.add_run("$a+b=c$")
    control.add_run(" 用于保留 inline math")

    doc.save(str(output_path))
    return output_path


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate minimal D40 mutant docx")
    parser.add_argument("--output", default=str(DEFAULT_OUTPUT), help="output .docx path")
    args = parser.parse_args()
    out = build_docx(Path(args.output))
    print(f"wrote {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
