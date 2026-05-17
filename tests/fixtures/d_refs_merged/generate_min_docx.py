"""Generate a minimal docx fixture with 5 custom-style refs paragraphs.

Path A: tests the integration path — verify pandoc emits 5 lines, not 1.
Currently passes under pandoc 3.9; serves as regression guard if pandoc upgrades
ever revert to the CASE-A merge-into-single-Para behavior.

Usage:
    python generate_min_docx.py [--output PATH]
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


DEFAULT_OUTPUT = Path(__file__).with_name("refs_merged_min.docx")
REFS_STYLE_NAME = "ReferencesEntry"  # custom paragraph style for refs entries


def build_docx(output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # Custom paragraph style for refs entries
    from docx.enum.style import WD_STYLE_TYPE
    styles = doc.styles
    if REFS_STYLE_NAME not in [s.name for s in styles]:
        styles.add_style(REFS_STYLE_NAME, WD_STYLE_TYPE.PARAGRAPH)

    # One chapter heading + body with [N] cites
    chap = doc.add_paragraph("第一章 引言")
    chap.style = "Heading 1"
    body = doc.add_paragraph()
    body.add_run("正文示例引用 [1][2][3][4][5].")

    # References section
    ref_head = doc.add_paragraph("参考文献")
    ref_head.style = "Heading 1"

    # 5 refs paragraphs, each with custom style
    entries = [
        "[1] Smith, J. A. (2020). Reference one title. Journal of X, 12(3), 45-67.",
        "[2] Doe, J. (2021). Reference two title. Book Publisher, Beijing.",
        "[3] Lee, K. (2019). Reference three title [J]. Journal Y, 8(2), 23-30.",
        "[4] Wang, M. (2022). Reference four title [D]. PhD thesis, University Z.",
        "[5] Chen, L. (2023). Reference five title [EB/OL]. https://example.com.",
    ]
    for entry in entries:
        para = doc.add_paragraph(entry)
        para.style = REFS_STYLE_NAME

    doc.save(str(output_path))
    return output_path


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate refs_merged minimal docx fixture")
    parser.add_argument("--output", default=str(DEFAULT_OUTPUT), help="output .docx path")
    args = parser.parse_args()
    out = build_docx(Path(args.output))
    print(f"wrote {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
