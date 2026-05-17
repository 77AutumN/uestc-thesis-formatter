"""Generate minimal docx fixtures for W2 docx_surgery tests.

Two fixtures:
  - surgery_relabel_minimal.docx: 段落用自定义 "1-1级" pStyle, 期望 relabel → "Heading 1"
  - surgery_inject_minimal.docx: body 含 "第一章 引言" 但 style 是 Normal, 期望 inject heading

Both use python-docx for fidelity (real Word OOXML 结构).

Usage: `python build_surgery_fixtures.py` regenerates both .docx files.
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent


def _set_pstyle(p, style_id: str):
    pPr = p._p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p._p.insert(0, pPr)
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        pStyle = OxmlElement("w:pStyle")
        pPr.insert(0, pStyle)
    pStyle.set(qn("w:val"), style_id)


def build_relabel():
    """custom 1-1级 → Heading 1 路径 (case_anon 模式)."""
    out = HERE / "surgery_relabel_minimal.docx"
    if out.exists():
        out.unlink()
    d = Document()
    # 注册 1-1级 (custom heading style)
    s1 = d.styles.add_style("1-1级", WD_STYLE_TYPE.PARAGRAPH)
    s1.base_style = d.styles["Normal"]
    s2 = d.styles.add_style("2-2级", WD_STYLE_TYPE.PARAGRAPH)
    s2.base_style = d.styles["Normal"]
    # 添加段落
    p1 = d.add_paragraph("第一章 引言")
    _set_pstyle(p1, s1.style_id)
    p2 = d.add_paragraph("1.1 研究背景")
    _set_pstyle(p2, s2.style_id)
    d.add_paragraph("正文段落.")
    p3 = d.add_paragraph("第二章 方法")
    _set_pstyle(p3, s1.style_id)
    d.save(out)
    print(f"wrote {out} ({out.stat().st_size} bytes)")


def build_inject():
    """body 含 "第一章" 但 style=Normal, 期望 inject_heading_before (case_anon 模式)."""
    out = HERE / "surgery_inject_minimal.docx"
    if out.exists():
        out.unlink()
    d = Document()
    # ToC 段落 (有标题但无 H1 style)
    d.add_paragraph("目录")
    d.add_paragraph("第一章 引言   1")
    d.add_paragraph("第二章 方法   3")
    # body 段落 — 文本含 "第N章" 但 style=Normal
    d.add_paragraph("第一章 引言")  # 这段需要 inject heading 或 promote 它
    d.add_paragraph("引言正文.")
    d.add_paragraph("第二章 方法")
    d.add_paragraph("方法正文.")
    d.save(out)
    print(f"wrote {out} ({out.stat().st_size} bytes)")


if __name__ == "__main__":
    build_relabel()
    build_inject()
