#!/usr/bin/env python3
"""
preflight_check.py — Word 文档预检模块

在 pipeline 执行提取之前，校验输入 .docx 是否满足 UESTC 论文模板要求。
失败项会中止 pipeline 并输出诊断报告，防止 garbage-in-garbage-out。

Usage:
    python preflight_check.py --input thesis.docx [--profile uestc-marxism]
"""

import argparse
import json
import os
import re
import sys


class PreflightReport:
    """预检报告收集器"""

    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.checks = []
        self.passed = 0
        self.failed = 0
        self.warnings = 0

    def check(self, name: str, passed: bool, detail: str = "", severity: str = "ERROR"):
        """记录一项检查结果"""
        status = "PASS" if passed else severity
        self.checks.append({
            "name": name,
            "status": status,
            "detail": detail
        })
        if passed:
            self.passed += 1
        elif severity == "ERROR":
            self.failed += 1
        else:
            self.warnings += 1

    @property
    def ok(self) -> bool:
        return self.failed == 0

    def summary(self) -> str:
        lines = [
            f"\n{'='*60}",
            f"  📋 Pre-flight 检查报告",
            f"{'='*60}",
            f"  文件: {os.path.basename(self.docx_path)}",
            f"  结果: {'✅ 全部通过' if self.ok else '❌ 存在阻塞项'}",
            f"  通过: {self.passed}  失败: {self.failed}  警告: {self.warnings}",
            f"{'='*60}",
        ]
        for c in self.checks:
            icon = "✅" if c["status"] == "PASS" else ("❌" if c["status"] == "ERROR" else "⚠️")
            lines.append(f"  {icon} {c['name']}")
            if c["detail"]:
                lines.append(f"     → {c['detail']}")
        lines.append(f"{'='*60}\n")
        return "\n".join(lines)

    def to_dict(self) -> dict:
        return {
            "docx_path": self.docx_path,
            "passed": self.passed,
            "failed": self.failed,
            "warnings": self.warnings,
            "ok": self.ok,
            "checks": self.checks
        }


def run_preflight(docx_path: str, profile: str = "uestc") -> PreflightReport:
    """执行所有预检项"""
    from docx import Document

    report = PreflightReport(docx_path)

    # === Check 0: 文件存在性 ===
    if not os.path.exists(docx_path):
        report.check("文件存在", False, f"文件不存在: {docx_path}")
        return report
    report.check("文件存在", True)

    # === Check 1: 文件大小 ===
    file_size = os.path.getsize(docx_path)
    if file_size < 10000:  # < 10KB 几乎肯定不是论文
        report.check("文件大小", False, f"文件过小 ({file_size} bytes)，可能不是有效论文")
    elif file_size > 100_000_000:  # > 100MB
        report.check("文件大小", False, f"文件过大 ({file_size // 1024 // 1024} MB)，请检查是否包含过多图片", "WARN")
    else:
        report.check("文件大小", True, f"{file_size // 1024} KB")

    # === Check 2: 文档可打开 ===
    try:
        doc = Document(docx_path)
    except Exception as e:
        report.check("文档可解析", False, f"python-docx 无法打开: {e}")
        return report
    report.check("文档可解析", True)

    # === Check 3: 段落数量 ===
    para_count = len(doc.paragraphs)
    if para_count < 50:
        report.check("段落数量", False, f"仅 {para_count} 个段落，论文通常 >200", "WARN")
    else:
        report.check("段落数量", True, f"{para_count} 个段落")

    # === Check 4: 封面表格结构 ===
    table_count = len(doc.tables)
    if table_count < 3:
        report.check("封面表格", False,
                      f"仅发现 {table_count} 个表格，UESTC 模板封面需要至少 3 个表格（中文封面/第二封面/英文封面）。"
                      "封面元数据将无法自动提取。", "WARN")
    else:
        # 检查 Table 0 是否包含封面关键字段
        t0 = doc.tables[0]
        t0_text = " ".join(c.text for row in t0.rows for c in row.cells)
        has_title_label = "论文题目" in t0_text
        has_author_label = "作者姓名" in t0_text
        has_advisor_label = "指导教师" in t0_text or "指导老师" in t0_text

        if has_title_label and has_author_label and has_advisor_label:
            report.check("封面表格", True, f"{table_count} 个表格，标签齐全（题目/作者/导师）")
        else:
            missing = []
            if not has_title_label: missing.append("论文题目")
            if not has_author_label: missing.append("作者姓名")
            if not has_advisor_label: missing.append("指导教师")
            report.check("封面表格", False,
                          f"封面表格缺少标签: {', '.join(missing)}", "WARN")

    # === Check 5: 章节标题 (Heading 1) ===
    heading1_paras = [p for p in doc.paragraphs
                      if p.style and p.style.name == 'Heading 1' and p.text.strip()]
    # 也检查手动编号的章节标题
    chapter_pattern = re.compile(r'^第[一二三四五六七八九十\d]+章\s')
    manual_chapters = [p for p in doc.paragraphs
                       if chapter_pattern.match(p.text.strip())]
    total_chapters = max(len(heading1_paras), len(manual_chapters))

    if total_chapters == 0:
        report.check("章节标题", False,
                      "未检测到任何 Heading 1 样式或「第X章」格式的段落。"
                      "extract_docx.py 将无法识别章节结构。")
    elif total_chapters < 3:
        report.check("章节标题", False,
                      f"仅发现 {total_chapters} 个章节，硕士论文通常 ≥5 章", "WARN")
    else:
        report.check("章节标题", True, f"{total_chapters} 个章节")

    # === Check 6: 中文摘要 ===
    all_text = " ".join(p.text for p in doc.paragraphs)
    has_abstract_zh = "摘" in all_text and "要" in all_text
    if has_abstract_zh:
        report.check("中文摘要", True)
    else:
        report.check("中文摘要", False, "未检测到「摘要」或「摘 要」关键词", "WARN")

    # === Check 7: 英文摘要 ===
    has_abstract_en = "ABSTRACT" in all_text.upper()
    if has_abstract_en:
        report.check("英文摘要", True)
    else:
        report.check("英文摘要", False, "未检测到「ABSTRACT」关键词", "WARN")

    # === Check 8: 参考文献 ===
    # 注意：部分论文模板用空格分隔写法：「参 考 文 献」
    all_text_nospace = all_text.replace(" ", "").replace("\u3000", "")
    has_refs = "参考文献" in all_text_nospace
    if has_refs:
        report.check("参考文献", True)
    else:
        report.check("参考文献", False, "未检测到「参考文献」或「参 考 文 献」段落")

    # === Check 9: 致谢 ===
    has_ack = "致" in all_text and "谢" in all_text
    if has_ack:
        report.check("致谢", True)
    else:
        report.check("致谢", False, "未检测到「致谢」段落", "WARN")

    # === Check 10: 引用标记 ===
    cite_markers = re.findall(r'\[\d+\]', all_text)
    footnote_markers = re.findall(r'[①②③④⑤⑥⑦⑧⑨⑩]', all_text)
    total_cites = len(cite_markers) + len(footnote_markers)
    if total_cites > 0:
        report.check("引用标记", True,
                      f"发现 {len(cite_markers)} 个方括号引用 + {len(footnote_markers)} 个圈号引用")
    else:
        report.check("引用标记", False, "未检测到任何引用标记 ([N] 或 ①②③)", "WARN")

    return report


def main():
    parser = argparse.ArgumentParser(description='论文 Word 文档预检工具')
    parser.add_argument('--input', required=True, help='输入 .docx 文件路径')
    parser.add_argument('--profile', default='uestc', help='论文 profile')
    parser.add_argument('--output', help='输出 JSON 报告路径')
    args = parser.parse_args()

    report = run_preflight(args.input, args.profile)
    print(report.summary())

    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            json.dump(report.to_dict(), f, ensure_ascii=False, indent=2)
        print(f"  报告已保存: {args.output}")

    sys.exit(0 if report.ok else 1)


if __name__ == '__main__':
    main()
