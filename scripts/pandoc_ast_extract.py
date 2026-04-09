#!/usr/bin/env python3
"""
pandoc_ast_extract.py — Pandoc AST 驱动的论文内容提取引擎

基于 Pandoc JSON AST 遍历实现文本提取和章节切分，替代旧的 python-docx 直接解析。
保持与 extract_docx.py 完全一致的输出格式，作为 drop-in 替代。

核心改进：
  - 混合标题检测：AST Header + "第X章" 正则兜底（解决 Word 样式不规范问题）
  - AST 原生文本提取：比 python-docx 段落遍历更精确
  - 为 Phase 2 (图片/公式/表格) 预留扩展点

Usage:
    python pandoc_ast_extract.py --input thesis.docx --output-dir ./output/extracted/
"""

import argparse
import json
import os
import re
import subprocess
import sys
import unicodedata


# ============================================================
# 1. Pandoc 调用层
# ============================================================

def run_pandoc(docx_path: str, media_dir: str = None) -> dict:
    """调用 pandoc 将 .docx 转为 JSON AST。

    Args:
        docx_path: .docx 文件路径
        media_dir: 可选，图片提取目录（Phase 2 用）

    Returns:
        Pandoc AST dict (包含 pandoc-api-version, meta, blocks)
    """
    cmd = ["pandoc", "-f", "docx", "-t", "json", docx_path]
    if media_dir:
        os.makedirs(media_dir, exist_ok=True)
        cmd.extend(["--extract-media", media_dir])

    try:
        result = subprocess.run(
            cmd, capture_output=True, text=True, encoding="utf-8", errors="replace",
            timeout=60
        )
        if result.returncode != 0:
            print(f"❌ Pandoc 运行失败 (exit {result.returncode}): {result.stderr[:300]}")
            sys.exit(1)
        return json.loads(result.stdout)
    except FileNotFoundError:
        print("❌ Pandoc 未安装或不在 PATH 中。请运行: winget install pandoc")
        sys.exit(1)
    except subprocess.TimeoutExpired:
        print("❌ Pandoc 超时 (60s)，文件可能过大")
        sys.exit(1)


# ============================================================
# 2. AST 文本提取工具
# ============================================================

def inlines_to_text(inlines: list) -> str:
    """将 AST inline 节点列表递归转为纯文本。

    处理的节点类型：
      Str, Space, SoftBreak, LineBreak, Emph, Strong,
      Superscript, Subscript, Strikeout, SmallCaps,
      Quoted, Span, Link, Note, Cite
    """
    parts = []
    for node in inlines:
        t = node.get("t", "")
        c = node.get("c")

        if t == "Str":
            parts.append(c)
        elif t in ("Space", "SoftBreak"):
            parts.append(" ")
        elif t == "LineBreak":
            parts.append("\n")
        elif t in ("Emph", "Strong", "Strikeout", "SmallCaps", "Superscript", "Subscript"):
            # 容器型 inline，递归提取子节点文本
            parts.append(inlines_to_text(c))
        elif t == "Quoted":
            # c = [quote_type, inlines]
            parts.append(inlines_to_text(c[1]))
        elif t == "Span":
            # c = [attrs, inlines]
            parts.append(inlines_to_text(c[1]))
        elif t == "Link":
            # c = [attrs, inlines, target]
            parts.append(inlines_to_text(c[1]))
        elif t == "Cite":
            # c = [citations, inlines]
            parts.append(inlines_to_text(c[1]))
        elif t == "Note":
            # 脚注，跳过（不混入正文）
            pass
        elif t == "Math":
            # c = [math_type, tex_string]
            parts.append(f"${c[1]}$")
        elif t == "Code":
            # c = [attrs, code_string]
            parts.append(c[1])
        elif t == "RawInline":
            # c = [format, raw_string]
            parts.append(c[1])
        elif t == "Image":
            # 跳过图片（不混入纯文本）
            pass
        # 其他未知类型静默忽略

    return "".join(parts)


def inlines_to_latex(inlines: list) -> str:
    """将 AST inline 节点列表转为 LaTeX 字符串。

    与 inlines_to_text 的区别：
      - Math 节点保留原始 LaTeX（不加 $ 包裹，交给调用方决定上下文）
      - 非 Math 文本做 LaTeX 转义
      - InlineMath → $...$
      - DisplayMath → \\[ ... \\]
    """
    parts = []
    for node in inlines:
        t = node.get("t", "")
        c = node.get("c")

        if t == "Str":
            parts.append(escape_latex(c))
        elif t in ("Space", "SoftBreak"):
            parts.append(" ")
        elif t == "LineBreak":
            parts.append("\n")
        elif t in ("Emph", "Strong", "Strikeout", "SmallCaps", "Superscript", "Subscript"):
            parts.append(inlines_to_latex(c))
        elif t == "Quoted":
            parts.append(inlines_to_latex(c[1]))
        elif t == "Span":
            parts.append(inlines_to_latex(c[1]))
        elif t == "Link":
            parts.append(inlines_to_latex(c[1]))
        elif t == "Cite":
            parts.append(inlines_to_latex(c[1]))
        elif t == "Note":
            # Note content is [Block] — typically [Para([inlines])]
            # Recursively extract text and wrap in \footnote{}
            note_parts = []
            for block in c:
                bt = block.get("t", "")
                if bt in ("Para", "Plain"):
                    note_parts.append(inlines_to_latex(block["c"]))
                elif bt == "OrderedList":
                    for item_blocks in block["c"][1]:
                        for ib in item_blocks:
                            if ib.get("t") in ("Para", "Plain"):
                                note_parts.append(inlines_to_latex(ib["c"]))
                elif bt == "BulletList":
                    for item_blocks in block["c"]:
                        for ib in item_blocks:
                            if ib.get("t") in ("Para", "Plain"):
                                note_parts.append(inlines_to_latex(ib["c"]))
            note_text = " ".join(note_parts).strip()
            if note_text:
                parts.append(f"\\footnote{{{note_text}}}")
        elif t == "Math":
            math_type = c[0].get("t", "InlineMath")
            tex = c[1]
            if math_type == "DisplayMath":
                parts.append(f"\\[ {tex} \\]")
            else:
                parts.append(f"${tex}$")
        elif t == "Code":
            parts.append(f"\\texttt{{{escape_latex(c[1])}}}")
        elif t == "RawInline":
            fmt, raw = c
            if fmt in ("latex", "tex"):
                parts.append(raw)
            else:
                parts.append(escape_latex(raw))
        elif t == "Image":
            pass  # Images handled at block level (Figure)

    return "".join(parts)


def block_to_text(block: dict) -> str:
    """将单个 AST block 节点转为纯文本。"""
    t = block.get("t", "")

    if t == "Para" or t == "Plain":
        return inlines_to_text(block["c"])
    elif t == "Header":
        level, attrs, inlines = block["c"]
        return inlines_to_text(inlines)
    elif t == "BlockQuote":
        return "\n\n".join(block_to_text(b) for b in block["c"])
    elif t == "OrderedList":
        items = block["c"][1]  # c = [list_attrs, [[blocks], ...]]
        texts = []
        for item_blocks in items:
            item_text = " ".join(block_to_text(b) for b in item_blocks)
            texts.append(item_text)
        return "\n".join(texts)
    elif t == "BulletList":
        texts = []
        for item_blocks in block["c"]:
            item_text = " ".join(block_to_text(b) for b in item_blocks)
            texts.append(item_text)
        return "\n".join(texts)
    elif t == "Div":
        # c = [attrs, blocks]
        return "\n\n".join(block_to_text(b) for b in block["c"][1])
    elif t == "RawBlock":
        return block["c"][1]
    elif t == "Table":
        # 简单提取 Table 内部的文字
        texts = []
        # Pandoc AST 遍历，找到所有的 Para/Plain 中的内容
        stack = [block.get("c", [])]
        while stack:
            n = stack.pop()
            if isinstance(n, list):
                stack.extend(reversed(n))
            elif isinstance(n, dict):
                if n.get("t") in ("Para", "Plain"):
                    texts.append(inlines_to_text(n.get("c", [])))
                elif "c" in n and (isinstance(n["c"], list) or isinstance(n["c"], dict)):
                    if isinstance(n["c"], list):
                        stack.extend(reversed(n["c"]))
                    else:
                        stack.append(n["c"])
        return "\n".join(texts)
    
    return ""


# ============================================================
# 3. 正则模式（与旧引擎一致）
# ============================================================

def normalize_text(text: str) -> str:
    """NFKC 全角字符规范化，并将换行符替换为空格（防标题跨行）"""
    return unicodedata.normalize("NFKC", text).replace("\n", " ").replace("\r", " ")


# 一级标题: "第X章"
RE_CHAPTER_CN = re.compile(r"^\s*第[一二三四五六七八九十百\d]+章\s+(.*)", re.UNICODE)

# 三级标题: "X.Y.Z" (必须在二级之前检查)
RE_SUBSECTION = re.compile(r"^\s*(\d+)\.(\d+)\.(\d+)\s*(.*)")

# 二级标题: "X.Y"
RE_SECTION = re.compile(r"^\s*(\d+)\.(\d+)\s*(.*)")


# 正文判定标点集合：包含这些标点的文本不是标题
BODY_PUNCTUATION = set("。；！？，、（）「」『』【】")


def is_body_text(text: str) -> bool:
    """判断文本是否为正文段落（而非标题）。

    规则（规范 §2.2：标题要突出重点、简明扼要，不要超过一行）：
      1. 长度 > 40 字符 → 正文
      2. 包含句号、分号等句子终结标点 → 正文
      3. 包含逗号且长度 > 20 → 正文（短标题如 "3.2 实验装置和方法" 不会触发）
    """
    stripped = text.strip()
    if len(stripped) > 40:
        return True
    if BODY_PUNCTUATION & set(stripped):
        return True
    # 含逗号且较长 — 可能是复合句
    if '，' in stripped and len(stripped) > 20:
        return True
    return False


def classify_paragraph(text: str):
    """对段落文本进行标题层级分类。

    Returns:
        (level_name, title_text) or None
        level_name: 'chapter', 'section', 'subsection', or None
        Returns None for body text (long text, sentence punctuation).
    """
    normalized = normalize_text(text)

    # 跳过目录条目：如果结尾是数字（页码），且前有空格，通常是目录泄漏
    if re.search(r"\s+\d+$", text.strip()) or re.search(r"\s+\d+$", normalized.strip()):
        return None

    # 章标题优先："第X章 XXX" 总是章标题，不受长度限制
    m = RE_CHAPTER_CN.match(normalized)
    if m:
        return ("chapter", m.group(1).strip())

    # 正文降级检测：在匹配 section/subsection 编号前先检查
    # 这防止 "3.4.1 我们使用SMW算法对这个矩阵求逆得到..." 被当 subsection
    if is_body_text(normalized):
        return None

    m = RE_SUBSECTION.match(normalized)
    if m:
        return ("subsection", m.group(4).strip() if m.group(4).strip() else normalized.strip())

    m = RE_SECTION.match(normalized)
    if m:
        return ("section", m.group(3).strip() if m.group(3).strip() else normalized.strip())

    return None


def escape_latex(text: str) -> str:
    """转义 LaTeX 特殊字符（保守策略）"""
    text = text.replace("&", "\\&")
    text = text.replace("%", "\\%")
    text = text.replace("#", "\\#")
    text = text.replace("_", "\\_")
    return text


# ============================================================
# 3.5 Figure 块处理 (Phase 2)
# ============================================================

# 图题编号正则："图2-1", "图 3-5", "图2.1" 等
RE_FIG_NUM = re.compile(r"^图\s*(\d+)[-.]?(\d+)\s*(.*)")

# 用于剥离 caption 中的冗余编号前缀
RE_CAPTION_NUM_PREFIX = re.compile(r"^图\s*\d+[-.]\d+\s*")

# 公式编号正则："(3-2)", "（2-1）" 等
RE_EQ_NUM = re.compile(r'[（(](\d+)[-.](\d+)[）)]')

# TOC 泄漏行正则："1.1 研究工作的背景与意义 1"
RE_TOC_LEAK = re.compile(r'^\d+(?:\.\d+){0,3}\s+\S.*\s+\d+\s*$')


def handle_figure_block(block: dict, media_base: str = "media") -> str:
    """将 Pandoc Figure block 转为 LaTeX figure 环境。

    Pandoc 3.x Figure 结构:
      Figure.c = [attrs, [caption_marker, caption_blocks], [body_blocks]]
      - caption_blocks: [Para [Str, Space, Str, ...]]
      - body_blocks:    [Plain [Image [attrs, alt, [src, title]]]]

    Args:
        block: AST Figure block dict
        media_base: 图片目录相对于编译根的路径

    Returns:
        LaTeX figure 环境字符串
    """
    c = block.get("c", [])
    if len(c) < 3:
        return "% [FIGURE: malformed block]\n"

    attrs, caption_data, body_blocks = c

    # --- 提取 caption ---
    caption_text = ""
    raw_caption = ""  # 保留原始 caption 用于 label 提取
    if caption_data and len(caption_data) >= 2 and caption_data[1]:
        # caption_data = [short_caption_or_null, [block, ...]]
        caption_blocks = caption_data[1]
        for cb in caption_blocks:
            if cb.get("t") in ("Para", "Plain"):
                raw_caption = inlines_to_text(cb["c"]).strip()
                # 剥离冗余编号前缀（LaTeX figure 环境自动编号）
                caption_text = RE_CAPTION_NUM_PREFIX.sub("", raw_caption).strip()
                break

    # --- 提取图片路径 ---
    img_src = ""
    img_width = ""
    for bb in body_blocks:
        if bb.get("t") in ("Para", "Plain"):
            for inline in bb.get("c", []):
                if inline.get("t") == "Image":
                    img_c = inline["c"]
                    # img_c = [attrs, alt_inlines, [src, title]]
                    img_attrs = img_c[0]  # ["", [], [["width","..."], ...]]
                    target = img_c[2]     # [src_path, title]
                    img_src = target[0] if target else ""
                    # 提取 width 属性
                    if len(img_attrs) >= 3:
                        for kv in img_attrs[2]:
                            if kv[0] == "width":
                                img_width = kv[1]
                    break

    if not img_src:
        return f"% [FIGURE: no image source, caption={caption_text}]\n"

    # --- 路径规范化：只保留 media/imageN.ext ---
    img_filename = os.path.basename(img_src)  # "image2.png"
    img_rel = f"{media_base}/{img_filename}"

    # --- LaTeX width：优先用原始宽度，最大 0.9\textwidth ---
    if img_width and "in" in img_width:
        try:
            width_in = float(img_width.replace("in", ""))
            # 估算：\textwidth ≈ 5.5in for UESTC template
            ratio = min(width_in / 5.5, 0.9)
            latex_width = f"{ratio:.2f}\\textwidth"
        except ValueError:
            latex_width = "0.8\\textwidth"
    else:
        latex_width = "0.8\\textwidth"

    # --- caption 处理 ---
    escaped_caption = escape_latex(caption_text) if caption_text else "图"

    # --- 生成 label（使用 raw_caption 匹配原始编号）---
    label = ""
    m = RE_FIG_NUM.match(raw_caption)
    if m:
        label = f"\\label{{fig:{m.group(1)}-{m.group(2)}}}"

    # --- 输出 ---
    lines = [
        "\n\\begin{figure}[H]",
        "  \\centering",
        f"  \\includegraphics[width={latex_width}]{{{img_rel}}}",
        f"  \\caption{{{escaped_caption}}}",
    ]
    if label:
        lines.append(f"  {label}")
    lines.append("\\end{figure}\n")
    return "\n".join(lines)


# ============================================================
# 3.6 Table 块处理 (Phase 2)
# ============================================================

_global_table_idx = 1
def handle_table_block(block: dict) -> str:
    """提取 Table 的纯文本内容并构造简化版 \begin{table}。
    增加复杂表格预警和 Override 占位符。
    """
    global _global_table_idx
    idx = _global_table_idx
    _global_table_idx += 1

    c = block.get("c", [])
    if len(c) < 5:
        return f"\n% [TABLE FALLBACK]\n{escape_latex(block_to_text(block))}\n"

    try:
        if len(c) == 6:
            attrs, caption, colspec, head, bodies, foot = c
        else:
            return f"\n% [TABLE FALLBACK OLD AST]\n{escape_latex(block_to_text(block))}\n"

        def extract_caption(cap):
            if len(cap) > 1 and cap[1]:
                return escape_latex(block_to_text(cap[1][0]))
            return ""
        cap_text = extract_caption(caption)

        is_complex = False

        def iter_rows(section):
            if not section: return []
            if len(section) == 2:
                rows = section[1]
            elif len(section) == 4:
                rows = section[2] + section[3]
            else:
                return []
            return [row[1] for row in rows if len(row) >= 2]

        def cell_to_latex(cell):
            nonlocal is_complex
            rowspan, colspan = 1, 1
            if len(cell) >= 5:
                rowspan, colspan = cell[2], cell[3]
                blocks = cell[4]
            elif len(cell) >= 2:
                blocks = cell[1]
            else:
                blocks = []

            if rowspan > 1 or colspan > 1:
                is_complex = True

            texts = []
            for b in blocks:
                if b.get("t") in ("Para", "Plain"):
                    texts.append(inlines_to_latex(b.get("c", [])).strip())
                else:
                    texts.append(escape_latex(block_to_text(b).strip()))
            
            content = " ".join(texts).replace("\n", " ").strip()
            
            # Retrieve col_count from the enclosing scope (it is evaluated later, but since cell_to_latex is called insideiter_rows after col_count is defined, we can rely on it if we bind it or it resolves). 
            # Actually, col_count is defined right after this function in the original code, so it's captured in the closure correctly!
            if colspan > 1:
                # Use m-columns for multicolumn to prevent text cutoff
                span_width = f"\\dimexpr 0.95\\textwidth * {colspan} / {max(1, col_count)} - 2\\tabcolsep\\relax"
                content = f"\\multicolumn{{{colspan}}}{{>{{\\centering\\arraybackslash}}m{{{span_width}}}}}{{{content}}}"
            if rowspan > 1:
                content = f"\\multirow{{{rowspan}}}{{*}}{{{content}}}"
            return content

        col_count = len(colspec) if colspec else 1
        
        res = f"\n% [TABLE-{idx}]\n"
        if is_complex:
            res += "% %% TODO: 复杂合并表格（RowSpan/ColSpan > 1），请人工复核 \\multicolumn 和 \\multirow\n"
            
        # Use longtable with wrapped centered cells (m-column from array package) to prevent horizontal overflow and allow page breaks
        col_type = f">{{\\centering\\arraybackslash}}m{{\\dimexpr 0.95\\textwidth/{col_count} - 2\\tabcolsep\\relax}}"
        res += f"\\begin{{longtable}}{{ *{{{col_count}}}{{{col_type}}} }}\n"
        
        if cap_text:
            res += f"  \\caption{{{cap_text}}} \\\\\n"
            
        res += "  \\toprule\n"
        
        # End of first head
        res += "  \\endfirsthead\n"
        if cap_text:
            # Continued caption for next pages (规范要求续表需要标注上接某表或续表)
            res += f"  \\caption*{{{cap_text} (续)}} \\\\\n"
        res += "  \\toprule\n"
        res += "  \\endhead\n"
        res += "  \\bottomrule\n"
        res += "  \\endfoot\n"
        res += "  \\bottomrule\n"
        res += "  \\endlastfoot\n"
        
        if head:
            for row in iter_rows(head):
                res += "  " + " & ".join(cell_to_latex(cell) for cell in row) + " \\\\\n"
            res += "  \\midrule\n"
            
        if bodies:
            for body in bodies:
                for row in iter_rows(body):
                    res += "  " + " & ".join(cell_to_latex(cell) for cell in row) + " \\\\\n"
                    
        if foot:
            res += "  \\midrule\n"
            for row in iter_rows(foot):
                res += "  " + " & ".join(cell_to_latex(cell) for cell in row) + " \\\\\n"
                
        res += "\\end{longtable}\n"
        res += f"% [/TABLE-{idx}]\n"
        
        return res
    except Exception as e:
        return f"\n% [TABLE EXCEPTION FALLBACK]\n{escape_latex(block_to_text(block))}\n"


# ============================================================
# 4. 混合章节检测
# ============================================================

# 章号提取正则
RE_CHAPTER_NUM = re.compile(r"第([一二三四五六七八九十百\d]+)章")

# 中文数字转阿拉伯数字映射
_CN_NUMS = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
            "六": 6, "七": 7, "八": 8, "九": 9, "十": 10}

def _cn_to_int(cn: str) -> int:
    """将中文数字转为 int（简化版，覆盖一~十）"""
    if cn.isdigit():
        return int(cn)
    return _CN_NUMS.get(cn, 0)


def find_chapters(blocks: list) -> list:
    """混合检测章节边界。

    策略:
      1. 遍历所有 Header L1 和 Para 块，用 "第X章" 正则匹配
      2. 按章号去重：目录(ToC)中的条目（带页码、block index 小）
         和正文中的标题可能产生重复，保留 block index 最大的那个
      3. 排除非章节的 Header L1（如 "攻读硕士学位期间取得的成果"）

    Returns:
        sorted list of dicts: {idx, raw_title, latex_title, filename, source}
    """
    # 第一遍：收集所有匹配，按章号分组
    by_chapter_num = {}  # chapter_num_str -> list of candidates

    for i, block in enumerate(blocks):
        t = block.get("t", "")
        text = ""

        if t == "Header":
            level, attrs, inlines = block["c"]
            text = inlines_to_text(inlines).strip()
            if level != 1:
                continue
        elif t == "Para":
            text = inlines_to_text(block["c"]).strip()
        elif t == "OrderedList":
            try:
                first_item_blocks = block["c"][1][0]
                if first_item_blocks and first_item_blocks[0]["t"] == "Para":
                    text = inlines_to_text(first_item_blocks[0]["c"]).strip()
                else:
                    continue
            except Exception:
                continue
        else:
            continue

        normalized = normalize_text(text)
        
        # NOTE: Some chapters may lack the "第X章" prefix in the source document.
        # If your document has this issue, add a case-specific normalization here.
        # Example: if "<chapter title>" in normalized and not normalized.startswith("第三章"):
        #     normalized = "第三章 " + normalized.replace("3.", "").strip()
        #     text = normalized
        
        # 跳过目录泄漏：如果结尾是数字，大概率是目录
        if re.search(r"\s+\d+$", normalized.strip()):
            continue

        m = RE_CHAPTER_CN.match(normalized)
        if m:
            # 提取章号
            num_match = RE_CHAPTER_NUM.search(normalized)
            if not num_match:
                continue
            ch_num_str = num_match.group(1)

            chapter_title = m.group(1).strip()
            # 去掉标题中可能带的页码（目录条目 "第一章 绪论 6"）
            raw_clean = re.sub(r"^第[一二三四五六七八九十百\d]+章\s*", "", normalized).strip()
            # 去掉末尾的纯数字（ToC 页码）
            raw_clean = re.sub(r"\s+\d+$", "", raw_clean).strip()

            candidate = {
                "idx": i,
                "raw_title": text,
                "latex_title": raw_clean if raw_clean else chapter_title,
                "ch_num_str": ch_num_str,
                "ch_num_int": _cn_to_int(ch_num_str),
                "source": "Header" if t == "Header" else "Para-regex",
            }

            if ch_num_str not in by_chapter_num:
                by_chapter_num[ch_num_str] = []
            by_chapter_num[ch_num_str].append(candidate)

    # 第二遍：按章号去重，每个章号只保留 block index 最大的（正文中的）
    deduped = []
    for ch_num_str, candidates in by_chapter_num.items():
        # 优先 Header，然后取 block index 最大的
        headers = [c for c in candidates if c["source"] == "Header"]
        best = headers[-1] if headers else candidates[-1]
        deduped.append(best)

    # 按章号排序
    deduped.sort(key=lambda c: c["ch_num_int"])

    # 分配文件名
    for i, ch in enumerate(deduped):
        ch["filename"] = f"ch{i + 1:02d}.tex"
        # 清理临时字段
        ch.pop("ch_num_str", None)
        ch.pop("ch_num_int", None)

    return deduped


# ============================================================
# 5. 特殊区段检测
# ============================================================

def find_special_sections(blocks: list, first_chapter_idx: int) -> dict:
    """检测摘要/致谢/参考文献/攻读成果等特殊区段。

    Returns:
        dict mapping section_name -> block_index
    """
    sections = {}

    for i, block in enumerate(blocks):
        t = block.get("t", "")
        text = ""

        if t == "Header":
            text = inlines_to_text(block["c"][2]).strip()
        elif t == "Para":
            text = inlines_to_text(block["c"]).strip()
        else:
            continue

        normalized = normalize_text(text).replace(" ", "").replace("\u3000", "")

        # 致谢（仅在正文之后检测，避免误匹配目录条目）
        if i > first_chapter_idx and "致" in normalized and "谢" in normalized and len(normalized) <= 10:
            sections["acknowledgement"] = i

        # 参考文献标题
        if i > first_chapter_idx and normalized == "参考文献":
            sections["references"] = i

        # 攻读成果
        if i > first_chapter_idx and "攻读" in normalized and "成果" in normalized:
            sections["accomplishments"] = i

    # 参考文献的特殊检测：在马院论文中，参考文献没有独立的 "参考文献" 标题块，
    # 而是致谢后面紧跟 OrderedList。需要检测这种模式。
    if "references" not in sections and "acknowledgement" in sections:
        ack_idx = sections["acknowledgement"]
        # 从致谢之后向后搜索 OrderedList（参考文献列表）
        for i in range(ack_idx + 1, min(ack_idx + 10, len(blocks))):
            if blocks[i].get("t") == "OrderedList":
                # 检查 OrderedList 中的第一项是否像参考文献
                items = blocks[i]["c"][1]
                if items:
                    first_item_text = ""
                    for sub in items[0]:
                        first_item_text += block_to_text(sub)
                    if "[M]" in first_item_text or "[J]" in first_item_text or "[D]" in first_item_text or "出版社" in first_item_text:
                        # 这是参考文献 OrderedList
                        sections["references_orderedlist"] = i
                        break

    # 摘要检测：基于位置（在封面表格之后、第一章之前）
    # 中文摘要 = 封面结束后第一段有意义的文本开始，到含 "关键词" 或 "Keywords" 的行
    # 英文摘要 = 中文摘要之后、目录/"目录" 之前
    abstract_zh_start = None
    abstract_en_start = None
    abstract_zh_kw_idx = None
    abstract_en_kw_idx = None

    for i, block in enumerate(blocks):
        if i >= first_chapter_idx:
            break

        t = block.get("t", "")
        if t == "Table":
            continue  # 跳过封面表格

        if t in ("Para", "Plain"):
            text = inlines_to_text(block["c"]).strip()
        elif t == "Header":
            # 一些 Word 文档将摘要内容导出为 Header blocks
            text = inlines_to_text(block["c"][2]).strip() if len(block["c"]) >= 3 else ""
        else:
            continue

        if not text:
            continue

        normalized = normalize_text(text)

            # 目录标记
        if normalized.replace(" ", "") == "目录":
            continue

            # 声明页检测（跳过）
        if any(kw in normalized for kw in (
                "本人声明", "作者签名", "导师签名", "独创性声明",
                "本学位论文作者", "涉密的学位论文", "使用授权",
            )):
            continue
        if re.match(r"^日期[：:]?\s*年\s*月\s*日$", normalized.replace(" ", "").replace("\u3000", "")):
            continue
            # 封面元数据行检测（跳过）
        if any(kw in normalized for kw in (
                "论文题目", "学科专业", "作者姓名", "指导老师", "指导教师",
                "学号", "学　号", "学  院", "申请学位", "培养单位",
                "答辩日期", "学位类别", "BACHELOR", "MASTER", "DOCTOR",
                "电子科技大学", "UNIVERSITY",
            )):
            continue
            # 跳过单独的标题行（摘要/ABSTRACT/独创性声明等加粗标题，通常很短且只有 Strong 节点）
        if t in ("Para", "Plain") and len(normalized) <= 10 and len(block["c"]) == 1 and block["c"][0].get("t") == "Strong":
            continue
            # 跳过 "摘要" / "摘 要" 标题独立行；同时用 ABSTRACT 触发英文摘要区域
        cleaned = normalized.replace(" ", "").replace("\u3000", "")
        if cleaned in ("摘要",):
            continue
        if cleaned == "ABSTRACT":
                # ABSTRACT 标题本身跳过，但标记下一段为英文摘要起始
            if abstract_zh_kw_idx and not abstract_en_start:
                    # 预标记：实际起始在下一个有意义的段落
                abstract_en_start = i + 1  # 临时，会在后续循环中被修正
            continue

            # 关键词行检测
        if "关键词" in normalized and not abstract_zh_kw_idx:
            abstract_zh_kw_idx = i
            if abstract_zh_start is None:
                abstract_zh_start = i  # 极端情况：关键词就是第一段
            continue

        if re.match(r"^Keywords?\s*[:：]", normalized, re.IGNORECASE) and not abstract_en_kw_idx:
            abstract_en_kw_idx = i
            if abstract_en_start is None:
                abstract_en_start = i
            continue

            # 根据语言检测确定摘要起始
        is_english = bool(re.match(r"^[A-Za-z]", normalized))

        if is_english and abstract_zh_kw_idx and not abstract_en_start:
                # 中文关键词之后出现的第一段英文 = 英文摘要起始
            abstract_en_start = i
        elif is_english and abstract_en_start and abstract_en_start > i:
                # 修正 ABSTRACT 标题预标记的起始位置
            abstract_en_start = i
        elif not is_english and not abstract_zh_start and abstract_zh_kw_idx is None:
                # 跳过封面后，第一段非声明的中文文本 = 中文摘要起始
            abstract_zh_start = i

    if abstract_zh_start is not None:
        sections["abstract_zh"] = abstract_zh_start
    if abstract_en_start is not None:
        sections["abstract_en"] = abstract_en_start

    return sections


# ============================================================
# 6. 引用标记检测与 cite_map 生成
# ============================================================

def detect_citation_markers(blocks: list, chapters: list, special: dict) -> int:
    """统计正文中 [数字] 引用标记数量"""
    # 正文范围：第一章到致谢/参考文献/文末
    if not chapters:
        return 0

    start = chapters[0]["idx"]
    end = special.get("acknowledgement",
           special.get("references",
           special.get("accomplishments", len(blocks))))

    count = 0
    for i in range(start, end):
        block = blocks[i]
        if block.get("t") in ("Para", "Plain"):
            text = inlines_to_text(block["c"])
            count += len(re.findall(r"\[\d+\]", text))
    return count


def generate_cite_map(references_raw: str, ref_count: int = 0) -> dict:
    """从参考文献原文生成 cite_map（编号 → BibTeX key）。

    支持两种格式：
      1. [数字] 开头的参考文献行（旧引擎格式）— 当大部分行匹配时使用
      2. 按行顺序编号（OrderedList 模式）— 当大部分行不匹配时使用
    """
    cite_map = {}

    non_empty_lines = [line.strip() for line in references_raw.split("\n") if line.strip()]
    total_lines = len(non_empty_lines)

    if total_lines == 0:
        return cite_map

    # 统计 [数字] 格式的行数
    bracketed_count = sum(1 for line in non_empty_lines if re.match(r"^\[\d+\]", line))

    # 如果大部分行都有 [num] 前缀（> 50%），使用编号模式
    if bracketed_count > total_lines * 0.5:
        for line in non_empty_lines:
            m = re.match(r"^\[(\d+)\]", line)
            if m:
                cite_map[m.group(1)] = f"ref{m.group(1)}"
    else:
        # 按行号顺序编号（OrderedList 模式或混合模式）
        for idx in range(1, total_lines + 1):
            cite_map[str(idx)] = f"ref{idx}"

    return cite_map


# ============================================================
# 7. LaTeX 章节生成
# ============================================================

def generate_chapter_tex(blocks: list, start_idx: int, end_idx: int,
                         latex_title: str, media_base: str = "media") -> str:
    """遍历 AST 区间，生成单章 .tex 内容。

    Phase 2 增强:
      - H1 汤修复：对 Header L1 做语义分类（section/subsection/body text）
      - Figure 块处理：Figure → \\begin{figure}[htbp]
      - Math 感知：Para 中的 Math inline 保留原始 LaTeX
      - BlockQuote 增强：像 Para 一样处理内含 Math 的引用块
    """
    tex_lines = [f"\\chapter{{{latex_title}}}\n"]
    last_figure_caption = ""  # W6: 用于 caption 去重

    for i in range(start_idx + 1, end_idx):
        block = blocks[i]
        t = block.get("t", "")

        if t == "Para":
            text = inlines_to_text(block["c"]).strip()
            if not text:
                continue

            # W6: Caption 去重 — 跳过与上一个 Figure caption 相同的段落
            if last_figure_caption and _text_similarity(text, last_figure_caption) > 0.7:
                last_figure_caption = ""  # 消费掉，只去重一次
                continue
            last_figure_caption = ""  # 非连续段落清除跟踪

            heading = classify_paragraph(text)
            if heading:
                level, title = heading
                if level == "chapter":
                    continue  # 跳过重复章标题
                elif level == "section":
                    tex_lines.append(f"\n\\section{{{escape_latex(title)}}}\n")
                elif level == "subsection":
                    tex_lines.append(f"\n\\subsection{{{escape_latex(title)}}}\n")
            else:
                # W2: 检测 DisplayMath + 编号模式 → equation 环境
                inlines = block["c"]
                has_display_math = any(
                    isinstance(n, dict) and n.get("t") == "Math"
                    and isinstance(n.get("c"), list) and len(n["c"]) >= 2
                    and isinstance(n["c"][0], dict) and n["c"][0].get("t") == "DisplayMath"
                    for n in inlines
                )
                if has_display_math:
                    # 提取公式 LaTeX 和可能的编号
                    math_tex = ""
                    for n in inlines:
                        if isinstance(n, dict) and n.get("t") == "Math":
                            if isinstance(n["c"][0], dict) and n["c"][0].get("t") == "DisplayMath":
                                math_tex = n["c"][1]
                                break
                    # 检查段落尾部是否有编号如 (3-2)
                    trailing = inlines_to_text(inlines[-5:]).strip() if len(inlines) > 3 else text
                    eq_match = RE_EQ_NUM.search(trailing)
                    if math_tex:
                        if eq_match:
                            ch_num, eq_num = eq_match.group(1), eq_match.group(2)
                            label = f"eq:{ch_num}-{eq_num}"
                            tex_lines.append(
                                f"\n\\begin{{equation}}\n"
                                f"  {math_tex}\n"
                                f"  \\label{{{label}}}\n"
                                f"\\end{{equation}}\n"
                            )
                        else:
                            # DisplayMath 无编号 → 使用 equation* (无编号居中)
                            tex_lines.append(
                                f"\n\\begin{{equation*}}\n"
                                f"  {math_tex}\n"
                                f"\\end{{equation*}}\n"
                            )
                    else:
                        # fallback: 原始逻辑
                        latex_text = inlines_to_latex(block["c"]).strip()
                        if latex_text:
                            tex_lines.append(f"\n{latex_text}\n")
                else:
                    # 使用 Math 感知的 LaTeX 输出
                    latex_text = inlines_to_latex(block["c"]).strip()
                    if latex_text:
                        tex_lines.append(f"\n{latex_text}\n")

        elif t == "Header":
            level, attrs, inlines = block["c"]
            text = inlines_to_text(inlines).strip()
            if not text:
                continue

            if level == 1:
                # H1 汤修复：用 classify_paragraph 做语义分类
                heading = classify_paragraph(text)
                if heading:
                    h_level, h_title = heading
                    if h_level == "chapter":
                        continue  # 跳过重复章标题
                    elif h_level == "section":
                        tex_lines.append(f"\n\\section{{{escape_latex(h_title)}}}\n")
                    elif h_level == "subsection":
                        tex_lines.append(f"\n\\subsection{{{escape_latex(h_title)}}}\n")
                else:
                    # H1 但不是标题 → 当作正文段落（STEM H1 汤）
                    latex_text = inlines_to_latex(inlines).strip()
                    if latex_text:
                        tex_lines.append(f"\n{latex_text}\n")
            elif level == 2:
                heading = classify_paragraph(text)
                has_math = any(n.get("t") == "Math" for n in inlines)
                if heading:
                    _, h_title = heading
                    tex_lines.append(f"\n\\section{{{escape_latex(h_title)}}}\n")
                elif has_math or "$" in text:
                    # H2 公式误标为标题 → 作为正文段落输出
                    latex_text = inlines_to_latex(inlines).strip()
                    if latex_text:
                        tex_lines.append(f"\n{latex_text}\n")
                elif is_body_text(text):
                    # H2 正文误标为标题 → 作为正文段落输出
                    latex_text = inlines_to_latex(inlines).strip()
                    if latex_text:
                        tex_lines.append(f"\n{latex_text}\n")
                else:
                    # H2 无编号、无 Math、无正文特征 → 仍作 section（短标题如 "ISAR"）
                    tex_lines.append(f"\n\\section{{{escape_latex(text)}}}\n")
            elif level == 3:
                heading = classify_paragraph(text)
                has_math = any(n.get("t") == "Math" for n in inlines)
                if heading:
                    _, h_title = heading
                    tex_lines.append(f"\n\\subsection{{{escape_latex(h_title)}}}\n")
                elif has_math or "$" in text:
                    # H3 公式误标为标题 → 作为正文段落输出
                    latex_text = inlines_to_latex(inlines).strip()
                    if latex_text:
                        tex_lines.append(f"\n{latex_text}\n")
                elif is_body_text(text):
                    # H3 正文误标为标题 → 作为正文段落输出
                    latex_text = inlines_to_latex(inlines).strip()
                    if latex_text:
                        tex_lines.append(f"\n{latex_text}\n")
                else:
                    # H3 无编号 → 降级为正文（规范§2.2 要求标题必须带编号）
                    latex_text = inlines_to_latex(inlines).strip()
                    if latex_text:
                        tex_lines.append(f"\n{latex_text}\n")
            elif level == 4:
                # 级联层级 4，使用加粗的段落模拟 \subsubsection
                latex_text = inlines_to_latex(inlines).strip()
                if latex_text:
                    tex_lines.append(f"\n\\vspace{{12pt}}\\noindent\\textbf{{{latex_text}}}\\vspace{{6pt}}\n")
            elif level >= 5:
                # 级联层级 5+ 等，完全视作正文
                latex_text = inlines_to_latex(inlines).strip()
                if latex_text:
                    tex_lines.append(f"\n{latex_text}\n")

        elif t == "Figure":
            # Phase 2: Figure → \begin{figure}[htbp]
            tex_lines.append(handle_figure_block(block, media_base=media_base))
            # W6: 记录 caption 用于后续段落去重
            c = block.get("c", [])
            if len(c) >= 3 and c[1] and len(c[1]) >= 2 and c[1][1]:
                for cb in c[1][1]:
                    if cb.get("t") in ("Para", "Plain"):
                        last_figure_caption = inlines_to_text(cb["c"]).strip()
                        last_figure_caption = RE_CAPTION_NUM_PREFIX.sub("", last_figure_caption).strip()
                        break

        elif t == "BlockQuote":
            # 增强：逐个子 block 处理，支持 Math 感知
            for sub_block in block.get("c", []):
                if sub_block.get("t") in ("Para", "Plain"):
                    latex_text = inlines_to_latex(sub_block["c"]).strip()
                    if latex_text:
                        tex_lines.append(f"\n{latex_text}\n")
                else:
                    sub_text = block_to_text(sub_block).strip()
                    if sub_text:
                        tex_lines.append(f"\n{escape_latex(sub_text)}\n")

        elif t == "OrderedList":
            # 中文论文自带编号（第一/第二/(1)/(2)），不使用 enumerate 避免缩进不一
            # 每个 sub-block 单独输出为一段，防止多段落 item 被合并
            items = block["c"][1]
            for item_blocks in items:
                for sub_block in item_blocks:
                    if sub_block.get("t") in ("Para", "Plain"):
                        latex_text = inlines_to_latex(sub_block["c"]).strip()
                        if latex_text:
                            tex_lines.append(f"\n{latex_text}\n")
                    else:
                        sub_text = block_to_text(sub_block).strip()
                        if sub_text:
                            tex_lines.append(f"\n{escape_latex(sub_text)}\n")

        elif t == "BulletList":
            # 同上，输出为普通段落保持一致缩进
            for item_blocks in block["c"]:
                for sub_block in item_blocks:
                    if sub_block.get("t") in ("Para", "Plain"):
                        latex_text = inlines_to_latex(sub_block["c"]).strip()
                        if latex_text:
                            tex_lines.append(f"\n{latex_text}\n")
                    else:
                        sub_text = block_to_text(sub_block).strip()
                        if sub_text:
                            tex_lines.append(f"\n{escape_latex(sub_text)}\n")

        elif t == "Table":
            # 引入加强版 handle_table_block
            tex_lines.append(handle_table_block(block))

        elif t == "RawBlock":
            # 保留原始 LaTeX/HTML 块
            fmt, raw = block["c"]
            if fmt == "latex" or fmt == "tex":
                tex_lines.append(f"\n{raw}\n")
                
        else:
            # C6 兜底暴力提取，宁死不丢文本
            fallback_text = block_to_text(block).strip()
            if fallback_text:
                tex_lines.append(f"\n% [BLOCK FALLBACK {t}]\n{escape_latex(fallback_text)}\n")

        # Div, HorizontalRule 等如果是已知但无需转换则可能被当作 fallback 输出

    final_output = "".join(tex_lines)
    # Post-process table notes (e.g., 资料来源：, 注：) for UESTC compliance
    final_output = re.sub(
        r'\n(资料来源[:：].*?)\n',
        r'\n{\\zihao{5}\\noindent \1 \\par}\n',
        final_output
    )
    final_output = re.sub(
        r'\n(注[:：].*?)\n',
        r'\n{\\zihao{5}\\noindent \1 \\par}\n',
        final_output
    )
    return final_output


# ============================================================
# 8. 文本块提取（摘要、致谢等）
# ============================================================

def is_toc_leak(text: str) -> bool:
    """检测文本是否为泄漏的 TOC 条目。
    
    匹配模式："1.1 研究工作的背景与意义 1" (编号 + 标题 + 页码)
    """
    lines = text.strip().split('\n')
    # 多行连续匹配 → 几乎确定是 TOC 块
    toc_count = sum(1 for l in lines if RE_TOC_LEAK.match(l.strip()))
    if toc_count >= 2:
        return True
    # 单行匹配 + 短文本 → 可能是 TOC
    if len(lines) == 1 and RE_TOC_LEAK.match(lines[0].strip()) and len(lines[0].strip()) < 80:
        return True
    return False


def _text_similarity(a: str, b: str) -> float:
    """简易文本相似度（基于字符集合交集比例）。用于 caption 去重。"""
    if not a or not b:
        return 0.0
    set_a, set_b = set(a.replace(' ', '')), set(b.replace(' ', ''))
    intersection = set_a & set_b
    return len(intersection) / max(len(set_a), len(set_b))


def extract_text_range(blocks: list, start_idx: int, end_idx: int) -> str:
    """提取指定 block 范围内的纯文本，自动过滤 TOC 泄漏行。"""
    lines = []
    for i in range(start_idx, end_idx):
        block = blocks[i]
        text = block_to_text(block).strip()
        if text:
            # TOC 泄漏过滤
            if is_toc_leak(text):
                continue
            lines.append(text)
    return "\n\n".join(lines)


# ============================================================
# 9. 封面元数据提取（复用旧 python-docx 逻辑）
# ============================================================

def extract_cover_metadata(docx_path: str) -> dict:
    """从 Word 封面表格提取元数据。直接复用旧引擎的 python-docx 逻辑。"""
    try:
        from docx import Document
    except ImportError:
        print("⚠️ python-docx 未安装，跳过封面元数据提取")
        return {}

    doc = Document(docx_path)
    meta = {}

    # === 中文封面：Table 0 ===
    if len(doc.tables) >= 1:
        t0 = doc.tables[0]
        for row in t0.rows:
            cells = [c.text.strip().replace("\u3000", " ") for c in row.cells]
            unique_cells = []
            seen = set()
            for c in cells:
                if c and c not in seen:
                    unique_cells.append(c)
                    seen.add(c)

            for i, cell_text in enumerate(unique_cells):
                if "论文题目" in cell_text and i + 1 < len(unique_cells):
                    meta["title_cn_part1"] = unique_cells[i + 1]
                elif "作者姓名" in cell_text and i + 1 < len(unique_cells):
                    meta["author_cn"] = unique_cells[i + 1].replace(" ", "")
                elif "指导教师" in cell_text and i + 1 < len(unique_cells):
                    meta["advisor_cn_raw"] = unique_cells[i + 1]
                elif "学科专业" in cell_text:
                    if i + 1 < len(unique_cells) and "学科" not in unique_cells[i + 1]:
                        meta["major_cn"] = unique_cells[i + 1]
                elif cell_text.replace(" ", "").replace("\u3000", "") == "学院":
                    if i + 1 < len(unique_cells):
                        meta["school_cn"] = unique_cells[i + 1]

        # 题目跨行处理
        title_parts = []
        found_title_label = False
        for row in t0.rows:
            cells_unique = list(dict.fromkeys(
                [c.text.strip() for c in row.cells if c.text.strip()]
            ))
            if any("论文题目" in c for c in cells_unique):
                found_title_label = True
                for c in cells_unique:
                    if "论文题目" not in c and c:
                        title_parts.append(c)
            elif found_title_label and title_parts:
                row_texts = list(dict.fromkeys(
                    [c.text.strip() for c in row.cells if c.text.strip()]
                ))
                if row_texts and not any(
                    k in row_texts[0] for k in ["学科", "学号", "作者", "指导", "学院"]
                ):
                    title_parts.extend(row_texts)
                found_title_label = False
            else:
                found_title_label = False
        if title_parts:
            meta["title_cn"] = "".join(title_parts)

    # === 英文封面：Table 2 ===
    if len(doc.tables) >= 3:
        t2 = doc.tables[2]
        for row in t2.rows:
            cells = [c.text.strip() for c in row.cells]
            unique = list(dict.fromkeys([c for c in cells if c]))
            if len(unique) >= 1:
                first = unique[0]
                if "Master Thesis" in first or "Submitted" in first:
                    continue
                if first == "Author" and len(unique) >= 2:
                    meta["author_en"] = unique[1]
                elif first == "Supervisor" and len(unique) >= 2:
                    meta["advisor_en"] = unique[1]
                elif first == "Discipline" and len(unique) >= 2:
                    meta["major_en"] = unique[1]
                elif first == "School" and len(unique) >= 2:
                    meta["school_en"] = unique[1]
                elif first == "Student ID" and len(unique) >= 2:
                    meta["student_id"] = unique[1]

        first_row = t2.rows[0]
        first_cells = list(dict.fromkeys(
            [c.text.strip() for c in first_row.cells if c.text.strip()]
        ))
        if first_cells and "Author" not in first_cells[0] and "Discipline" not in first_cells[0]:
            meta["title_en"] = first_cells[0]

    # === Table 1 补充 ===
    if len(doc.tables) >= 2 and "title_cn" not in meta:
        t1 = doc.tables[1]
        for row in t1.rows:
            cells_unique = list(dict.fromkeys(
                [c.text.strip() for c in row.cells if c.text.strip()]
            ))
            if len(cells_unique) == 1 and len(cells_unique[0]) > 10:
                text = cells_unique[0]
                if "题名" not in text and "注" not in text and "学位" not in text:
                    meta["title_cn"] = text
                    break

    # === 清理导师 ===
    if "advisor_cn_raw" in meta:
        raw = meta["advisor_cn_raw"]
        cleaned = re.sub(r"[\s\u3000]+", " ", raw).strip()
        parts = cleaned.split(" ")
        if len(parts) >= 2:
            meta["advisor_name_cn"] = parts[0]
            meta["advisor_title_cn"] = " ".join(parts[1:])
        else:
            meta["advisor_name_cn"] = cleaned
            meta["advisor_title_cn"] = ""
        del meta["advisor_cn_raw"]

    if "title_cn_part1" in meta and "title_cn" not in meta:
        meta["title_cn"] = meta["title_cn_part1"]
    meta.pop("title_cn_part1", None)

    return meta


# ============================================================
# 9.5 段落式封面元数据提取（STEM 论文兜底）
# ============================================================

# 封面字段匹配模式（Tab/空格分隔的 key-value 行）
_COVER_PATTERNS = [
    ("title_field", re.compile(r"^论文题目[\s\t]+(.*)", re.UNICODE)),
    ("school_cn",   re.compile(r"^学\s*院[\s\t]+(.*)", re.UNICODE)),
    ("major_cn",    re.compile(r"^专\s*业[\s\t]+(.*)", re.UNICODE)),
    ("student_id",  re.compile(r"^学\s*号[\s\t]+(\d+)", re.UNICODE)),
    ("author_cn",   re.compile(r"^作者姓名[\s\t]+(.*)", re.UNICODE)),
    ("advisor_cn_raw", re.compile(r"^指导(?:教师|老师)[\s\t]+(.*)", re.UNICODE)),
]

# 大学标识行（用于定位封面区域）
RE_UNIVERSITY_HEADER = re.compile(
    r"电\s*子\s*科\s*技\s*大\s*学|UNIVERSITY OF ELECTRONIC", re.IGNORECASE
)


def extract_cover_metadata_from_ast(blocks: list, first_chapter_idx: int) -> dict:
    """从 AST 段落中提取封面元数据（段落式排版，无 Word 表格）。

    适用于 STEM 论文，其封面通过 Tab/空格对齐字段，而非表格单元格。

    Args:
        blocks: Pandoc AST blocks
        first_chapter_idx: 第一章在 blocks 中的 index

    Returns:
        dict: 与 extract_cover_metadata() 输出格式一致的元数据
        同时附带 '_cover_block_indices' 记录匹配到的 block 索引
    """
    meta = {}
    cover_indices = []  # 记录封面块索引
    title_parts = []    # 标题可能跨行

    scan_end = min(first_chapter_idx, 50)  # 封面不会超过前 50 个 block

    def _scan_para(normalized, i, meta, title_parts, cover_indices):
        """内部辅助：对单个段落文本做封面字段匹配"""
        # 跳过大学标识行（不是元数据）
        if RE_UNIVERSITY_HEADER.search(normalized):
            cover_indices.append(i)
            return True

        # 匹配学位论文/THESIS 标识行，同时提取 degree_type
        _degree_match = re.match(r"^(学士|硕士|博士|工程)?学位论文$", normalized.replace(" ", ""))
        if _degree_match:
            cover_indices.append(i)
            _dw = _degree_match.group(1) or ""
            _degree_map = {"学士": "bachelor", "硕士": "master", "博士": "doctor", "工程": "engdoctor"}
            if _dw in _degree_map:
                meta["degree_type"] = _degree_map[_dw]
            return True
        _eng_degree_match = re.match(
            r"^(BACHELOR|MASTER|DOCTORAL?)\s*(THESIS|DISSERTATION)$",
            normalized, re.IGNORECASE,
        )
        if _eng_degree_match:
            cover_indices.append(i)
            _ed = _eng_degree_match.group(1).upper()
            _eng_map = {"BACHELOR": "bachelor", "MASTER": "master", "DOCTOR": "doctor", "DOCTORAL": "doctor"}
            if _ed in _eng_map:
                meta.setdefault("degree_type", _eng_map[_ed])
            return True

        # 匹配封面字段
        for field_name, pattern in _COVER_PATTERNS:
            m = pattern.match(normalized)
            if m:
                value = m.group(1).strip()
                if field_name == "title_field":
                    title_parts.append(value)
                else:
                    meta[field_name] = value
                cover_indices.append(i)
                return True

        # 标题跨行检测：紧跟"论文题目"行的短文本行
        if title_parts and not meta.get("school_cn"):
            if len(normalized) < 50 and not any(
                k in normalized
                for k in ("学", "专", "作者", "指导", "声明", "摘要")
            ):
                title_parts.append(normalized)
                cover_indices.append(i)
                return True

        return False

    for i in range(scan_end):
        block = blocks[i]
        bt = block.get("t")

        # 处理 BlockQuote（封面字段可能被 Word 转为引用块）
        if bt == "BlockQuote":
            for sub in block.get("c", []):
                if sub.get("t") in ("Para", "Plain"):
                    text = inlines_to_text(sub["c"]).strip()
                    if text:
                        normalized = normalize_text(text).replace("\u3000", " ")
                        _scan_para(normalized, i, meta, title_parts, cover_indices)
            continue

        if bt not in ("Para", "Plain"):
            continue
        text = inlines_to_text(block["c"]).strip()
        if not text:
            continue

        normalized = normalize_text(text).replace("\u3000", " ")
        _scan_para(normalized, i, meta, title_parts, cover_indices)

    # 组装标题
    if title_parts:
        meta["title_cn"] = "".join(title_parts)

    # 清理导师字段（与 extract_cover_metadata 一致）
    if "advisor_cn_raw" in meta:
        raw = meta["advisor_cn_raw"]
        cleaned = re.sub(r"[\s\u3000]+", " ", raw).strip()
        parts = cleaned.split(" ")
        if len(parts) >= 2:
            meta["advisor_name_cn"] = parts[0]
            meta["advisor_title_cn"] = " ".join(parts[1:])
        else:
            meta["advisor_name_cn"] = cleaned
            meta["advisor_title_cn"] = ""
        del meta["advisor_cn_raw"]

    # 记录封面块索引（供 strip 使用）
    meta["_cover_block_indices"] = sorted(set(cover_indices))

    return meta


def strip_cover_and_toc_blocks(blocks: list, cover_indices: list,
                                first_chapter_idx: int) -> int:
    """标记封面和 TOC 文本块为已处理，防止污染摘要和正文。

    不删除 blocks（保持索引稳定），将匹配块的 type 改为 'Null'。

    Args:
        blocks: Pandoc AST blocks（会被原地修改）
        cover_indices: 封面元数据块的索引列表
        first_chapter_idx: 第一章在 blocks 中的 index

    Returns:
        被置空的 block 数量
    """
    stripped = 0

    # 1. 标记封面块
    for idx in cover_indices:
        if 0 <= idx < len(blocks):
            blocks[idx]["t"] = "Null"
            stripped += 1

    # 2. 检测并标记 Word 生成的 TOC 文本块
    RE_TOC_ENTRY = re.compile(
        r"^(?:第[一二三四五六七八九十\d]+章|[\d]+\.[\d]+"
        r"|摘\s*要|ABSTRACT|致\s*谢|参考文献|目\s*录).*\d+\s*$",
        re.UNICODE,
    )
    toc_started = False

    for i in range(first_chapter_idx):
        block = blocks[i]
        if block.get("t") == "Null":
            continue
        if block.get("t") not in ("Para", "Plain"):
            continue

        text = inlines_to_text(block["c"]).strip()
        normalized = normalize_text(text).replace("\u3000", " ")

        # 检测 "目录" 标题行
        if normalized.replace(" ", "") == "目录":
            blocks[i]["t"] = "Null"
            toc_started = True
            stripped += 1
            continue

        # TOC 区域内的条目
        if toc_started and RE_TOC_ENTRY.match(normalized):
            blocks[i]["t"] = "Null"
            stripped += 1

    return stripped


# ============================================================
# 10. 主函数
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="Pandoc AST 驱动的论文内容提取引擎"
    )
    parser.add_argument("--input", required=True, help="输入 .docx 文件路径")
    parser.add_argument("--output-dir", required=True, help="输出目录")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"❌ 文件不存在: {args.input}")
        sys.exit(1)

    output_dir = args.output_dir
    chapters_dir = os.path.join(output_dir, "chapters")
    # 清理旧的章节文件（避免残留）
    if os.path.exists(chapters_dir):
        import shutil
        shutil.rmtree(chapters_dir)
    os.makedirs(chapters_dir, exist_ok=True)

    # === Step 1: Pandoc 解析（Phase 2: 启用图片提取）===
    media_dir = os.path.join(output_dir, "media")
    print(f"📖 [AST Engine] 正在解析: {args.input}")
    ast = run_pandoc(args.input, media_dir=media_dir)
    blocks = ast.get("blocks", [])
    print(f"  Pandoc API: {ast.get('pandoc-api-version', '?')}")
    print(f"  共 {len(blocks)} 个 AST blocks")
    # 统计 Figure 和 Math
    fig_count = sum(1 for b in blocks if b.get("t") == "Figure")
    math_count = 0
    for b in blocks:
        if b.get("t") in ("Para", "Header"):
            inl = b.get("c", [])
            if b["t"] == "Header":
                inl = b["c"][2] if len(b.get("c", [])) >= 3 else []
            for node in (inl if isinstance(inl, list) else []):
                if isinstance(node, dict) and node.get("t") == "Math":
                    math_count += 1
    print(f"  📊 Figure: {fig_count} | Math: {math_count}")

    # 统计 block 类型
    from collections import Counter
    type_counts = Counter(b["t"] for b in blocks)
    print(f"  类型分布: {dict(type_counts.most_common())}")

    # === Step 2: 混合章节检测 ===
    chapters = find_chapters(blocks)
    print(f"\n  📌 识别到 {len(chapters)} 个章节:")
    for ch in chapters:
        src = "✓ Header" if ch["source"] == "Header" else "⚡ Para-regex"
        print(f"    {ch['filename']}: {ch['raw_title']}  [{src}]")

    if len(chapters) == 0:
        print("❌ 未检测到任何章节标题！请检查文档结构。")
        sys.exit(1)

    # === Step 3: 特殊区段检测 ===
    first_ch_idx = chapters[0]["idx"] if chapters else len(blocks)
    special = find_special_sections(blocks, first_ch_idx)
    print(f"\n  📋 特殊区段: {list(special.keys())}")
    for name, idx in special.items():
        print(f"    {name}: block[{idx}]")

    # === Step 4: 引用标记检测 ===
    citation_count = detect_citation_markers(blocks, chapters, special)
    print(f"\n  🔗 正文引用标记 [数字]: {citation_count} 处")

    # === Step 5: 输出 outline.json ===
    outline = {
        "chapters": [
            {
                "filename": ch["filename"],
                "title": ch["raw_title"],
                "latex_title": ch["latex_title"],
            }
            for ch in chapters
        ],
        "special_sections": {k: True for k in special.keys()},
    }
    outline_path = os.path.join(output_dir, "outline.json")
    with open(outline_path, "w", encoding="utf-8") as f:
        json.dump(outline, f, ensure_ascii=False, indent=2)
    print(f"  ✅ outline.json")

    # === Step 6: thesis_meta.json ===
    meta = {
        "total_paragraphs": len([b for b in blocks if b["t"] in ("Para", "Plain")]),
        "total_chapters": len(chapters),
        "citation_markers_in_body": citation_count,
        "has_abstract_zh": "abstract_zh" in special,
        "has_abstract_en": "abstract_en" in special,
        "has_acknowledgement": "acknowledgement" in special,
        "has_references": "references" in special or "references_orderedlist" in special,
        "has_accomplishments": "accomplishments" in special,
        "extraction_engine": "pandoc-ast",
    }
    meta_path = os.path.join(output_dir, "thesis_meta.json")
    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)
    print(f"  ✅ thesis_meta.json")

    # === Step 6.5: 封面元数据（Fallback 链：表格 → AST 段落）===
    print(f"\n  📋 提取封面元数据...")
    cover_meta = extract_cover_metadata(args.input)
    if cover_meta:
        print(f"  ✅ 表格模式提取成功")
    else:
        print(f"  ⚠️ 表格模式未匹配，尝试 AST 段落模式...")
        cover_meta = extract_cover_metadata_from_ast(blocks, first_ch_idx)
        if cover_meta and any(k in cover_meta for k in ("title_cn", "author_cn")):
            print(f"  ✅ AST 段落模式提取成功")
        else:
            print(f"  ⚠️ 未能提取封面元数据（两种模式均未匹配）")

    # 分离内部索引信息，不写入 JSON
    cover_block_indices = cover_meta.pop("_cover_block_indices", [])

    if cover_meta:
        cover_path = os.path.join(output_dir, "cover_metadata.json")
        with open(cover_path, "w", encoding="utf-8") as f:
            json.dump(cover_meta, f, ensure_ascii=False, indent=2)
        print(f"  ✅ cover_metadata.json")
        for k, v in cover_meta.items():
            print(f"    {k}: {v}")

    # === Step 6.6: 封面/TOC 块剥离（防止元数据污染摘要和正文）===
    stripped_count = strip_cover_and_toc_blocks(blocks, cover_block_indices, first_ch_idx)
    if stripped_count > 0:
        print(f"  🧹 已标记 {stripped_count} 个封面/TOC 块为 Null（防泄漏）")

    # === 确定最末章节的结束边界 ===
    # 边界优先级：致谢 > 参考文献 > 参考文献OrderedList > 攻读成果 > 文末
    last_chapter_end = len(blocks)
    for key in ("acknowledgement", "references", "references_orderedlist", "accomplishments"):
        if key in special:
            last_chapter_end = special[key]
            break

    # === Step 7: 生成章节 .tex 文件 ===
    print(f"\n  📝 生成章节 .tex 文件:")
    for ch_idx, ch_info in enumerate(chapters):
        if ch_idx + 1 < len(chapters):
            end_idx = chapters[ch_idx + 1]["idx"]
        else:
            end_idx = last_chapter_end

        tex_content = generate_chapter_tex(
            blocks, ch_info["idx"], end_idx, ch_info["latex_title"],
            media_base="media"
        )
        filepath = os.path.join(chapters_dir, ch_info["filename"])
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(tex_content)
        # 统计段落数
        para_count = tex_content.count("\n\n")
        print(f"    ✅ {ch_info['filename']} ({para_count} 段)")

    # === Step 8: 致谢 ===
    if "acknowledgement" in special:
        ack_start = special["acknowledgement"]
        # 致谢结束 = 参考文献 OrderedList 或参考文献标题 或攻读成果 或文末
        ack_end = special.get(
            "references_orderedlist",
            special.get("references",
            special.get("accomplishments", len(blocks))),
        )
        # 跳过 "致 谢" 标题行本身
        ack_text = extract_text_range(blocks, ack_start + 1, ack_end)
        with open(os.path.join(output_dir, "acknowledgement.txt"), "w", encoding="utf-8") as f:
            f.write(ack_text)
        print(f"  ✅ acknowledgement.txt")

    # === Step 9: 参考文献 ===
    references_raw = ""
    if "references" in special:
        # 有独立的 "参考文献" 标题块
        ref_start = special["references"]
        ref_end = special.get("accomplishments", len(blocks))
        references_raw = extract_text_range(blocks, ref_start + 1, ref_end)
    elif "references_orderedlist" in special:
        # 参考文献在 OrderedList 中（马院论文模式）
        ref_list_idx = special["references_orderedlist"]
        ref_block = blocks[ref_list_idx]
        if ref_block["t"] == "OrderedList":
            items = ref_block["c"][1]
            ref_lines = []
            for item_blocks in items:
                item_text = " ".join(block_to_text(b) for b in item_blocks).strip()
                if item_text:
                    ref_lines.append(item_text)
            references_raw = "\n".join(ref_lines)
            print(f"  📚 从 OrderedList 提取 {len(ref_lines)} 条参考文献")

    if references_raw:
        with open(os.path.join(output_dir, "references_raw.txt"), "w", encoding="utf-8") as f:
            f.write(references_raw)
        print(f"  ✅ references_raw.txt")

    # === Step 9.5: cite_map.json ===
    if references_raw:
        cite_map = generate_cite_map(references_raw)
        cite_map_path = os.path.join(output_dir, "cite_map.json")
        with open(cite_map_path, "w", encoding="utf-8") as f:
            json.dump(cite_map, f, ensure_ascii=False, indent=2)
        print(f"  ✅ cite_map.json ({len(cite_map)} 条)")

    # === Step 10: 摘要 ===
    if "abstract_zh" in special:
        zh_start = special["abstract_zh"]
        # 中文摘要结束 = 英文摘要起始，或第一章
        zh_end = special.get("abstract_en", first_ch_idx)
        # 找 "关键词" 行的位置，摘要内容包含关键词
        abs_text = extract_text_range(blocks, zh_start, zh_end)
        with open(os.path.join(output_dir, "abstract_zh.txt"), "w", encoding="utf-8") as f:
            f.write(abs_text)
        print(f"  ✅ abstract_zh.txt")

    if "abstract_en" in special:
        en_start = special["abstract_en"]
        # 英文摘要结束 = "目录" 或第一章
        en_end = first_ch_idx
        # 查找 "目录" block 作为更精确的结束点
        for i in range(en_start, first_ch_idx):
            if blocks[i].get("t") == "Para":
                text = inlines_to_text(blocks[i]["c"]).strip()
                if normalize_text(text).replace(" ", "") == "目录":
                    en_end = i
                    break
        abs_text = extract_text_range(blocks, en_start, en_end)
        with open(os.path.join(output_dir, "abstract_en.txt"), "w", encoding="utf-8") as f:
            f.write(abs_text)
        print(f"  ✅ abstract_en.txt")

    # === Step 11: 攻读成果 ===
    if "accomplishments" in special:
        acc_start = special["accomplishments"]
        # 成果结束 = 文末（成果通常是最后一个 section）
        acc_end = len(blocks)
        # 跳过标题行本身
        acc_text = extract_text_range(blocks, acc_start + 1, acc_end)
        if acc_text.strip():
            with open(os.path.join(output_dir, "accomplishment.txt"), "w", encoding="utf-8") as f:
                f.write(acc_text)
            print(f"  ✅ accomplishment.txt")

    # === 完成 ===
    print(f"\n🎉 [AST Engine] 提取完成! 输出目录: {output_dir}")
    print(f"   章节: {len(chapters)} | 引用标记: {citation_count} | 引擎: pandoc-ast")


if __name__ == "__main__":
    main()
