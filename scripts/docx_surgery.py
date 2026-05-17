"""docx_surgery.py — Round 9 W2 B 类结构修复事务引擎 (v0.1.0).

CLI 三段:
  plan  : 跑所有 detector, 生成 pre_surgery_manifest.json + surgery_plan.json
  apply : 事务原子性: backup -> temp -> ops -> post manifest -> verify -> 替换 source
  verify: 独立基于 manifest delta, 不调 risk-router. 写 surgery_verify_report.json

operation type 4 个:
  - register_heading_style
  - relabel_pstyle
  - inject_heading_before  (W2 stub, 主体 W3 补)
  - delete_orphan_title_paragraph  (W2 stub, 主体 W3 补)

详见 reference/docx_surgery_plan.schema.md.
"""
from __future__ import annotations
import argparse
import hashlib
import json
import os
import re
import shutil
import sys
import tempfile
import zipfile
from datetime import datetime, timezone
from typing import Dict, List, Optional, Tuple

THIS_DIR = os.path.dirname(os.path.abspath(__file__))
if THIS_DIR not in sys.path:
    sys.path.insert(0, THIS_DIR)

import source_manifest as sm  # noqa: E402

PLAN_SCHEMA_VERSION = "0.1.0"


def _sha256_file(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def _now_iso() -> str:
    return datetime.now(timezone.utc).astimezone().isoformat()


def _visible_text_hash(docx_path: str) -> str:
    """计算 docx body 所有 <w:t> 文字 concat 后归一化 (压缩空白) 的 sha256.

    归一化避免 python-docx Document.save() 的非语义重序列化误报
    (例如相邻 <w:t> merge / 空白规范化).
    """
    with zipfile.ZipFile(docx_path) as z:
        doc = z.read("word/document.xml").decode("utf-8", errors="replace")
    body_match = re.search(r"<w:body>(.*?)</w:body>", doc, re.DOTALL)
    body = body_match.group(1) if body_match else doc
    texts = re.findall(r"<w:t[^>]*>([^<]*)</w:t>", body)
    full = "".join(texts)
    # 归一化 + decode XML entities (python-docx 重序列化会把 &quot; 转 ")
    full = (full.replace("&quot;", '"').replace("&amp;", "&")
            .replace("&lt;", "<").replace("&gt;", ">").replace("&apos;", "'"))
    # 压缩所有空白 (含全角空格/NBSP) 为单个 ASCII 空格, 去掉首尾
    full = re.sub(r"[\s 　]+", " ", full).strip()
    h = hashlib.sha256()
    h.update(full.encode("utf-8"))
    return h.hexdigest()


# ============================================================
# Detectors — 输入: probe manifest, 输出: List[operation dict]
# ============================================================

def _detector_register_heading_style(manifest: Dict) -> List[Dict]:
    """检测 styles.xml 缺 Heading 1/2/3 (实际通过看 paragraphs 是否有人用 'Heading N')."""
    ops: List[Dict] = []
    style_names = set()
    custom_levels = set()
    for p in manifest.get("paragraphs", []):
        sn = p.get("style_name", "")
        if sn:
            style_names.add(sn)
        if sn in {"1-1级", "2-2级", "3-3级"}:
            level = int(sn[0])
            custom_levels.add(level)
    for level in sorted(custom_levels):
        target_name = f"Heading {level}"
        if target_name in style_names:
            continue
        ops.append({
            "id": "",  # caller assigns
            "type": "register_heading_style",
            "enabled": True,
            "risk_class": "B",
            "visible_text_change": False,
            "params": {"level": level},
            "evidence": [f"style {target_name!r} missing in styles.xml; need to map custom heading"],
            "expected_delta": {"styles_added": [target_name]},
            "confidence": 1.0,
            "target_paragraph_ids": [],
        })
    return ops


def _detector_relabel_pstyle(manifest: Dict) -> List[Dict]:
    """检测 custom 样式段 (1-1级 etc) 该 relabel 为 Heading 1/2/3."""
    style_map = {"1-1级": "Heading 1", "2-2级": "Heading 2", "3-3级": "Heading 3"}
    target_paras: Dict[str, List[str]] = {tgt: [] for tgt in set(style_map.values())}
    counts: Dict[str, int] = {k: 0 for k in style_map}
    matched: List[str] = []
    for p in manifest.get("paragraphs", []):
        sn = p.get("style_name", "")
        if sn in style_map:
            target_paras[style_map[sn]].append(p["id"])
            counts[sn] += 1
            matched.append(p["id"])
    if not matched:
        return []
    return [{
        "id": "",
        "type": "relabel_pstyle",
        "enabled": True,
        "risk_class": "B",
        "visible_text_change": False,
        "params": {"style_map": style_map},
        "evidence": [
            f"{counts['1-1级']} paragraphs use '1-1级'",
            f"{counts['2-2级']} paragraphs use '2-2级'",
            f"{counts['3-3级']} paragraphs use '3-3级'",
        ],
        "expected_delta": {
            "heading_counts": {tgt: len(ids) for tgt, ids in target_paras.items() if ids},
        },
        "confidence": 0.95,
        "target_paragraph_ids": matched,
    }]


def _detector_inject_heading_before(manifest: Dict) -> List[Dict]:
    """检测 body 含"第N章"段 + 与 anchor 之前的章引言段配对 (W4-C 升级版).

    case14 类型 (Heading counts 0/0/0 + ToC 含 第N章 但 body 没相应 H1) 是基本场景.
    case19 round 2/2c 凝结的 anchor 上移 = 章引言段配对 (W4-C 主体).

    输出:
      - inject ops, anchor_text_match = 章引言段 text (若有), 否则 = 第N章段 text
      - visible_text_change=True (新增可见文字)
    """
    headings = manifest.get("headings", [])
    h1_count = sum(1 for h in headings if h.get("level") == 1 and h.get("status") == "accepted")
    custom_h1 = sum(1 for h in headings if h.get("level") == 1 and h.get("source") == "custom_style")
    if h1_count > 0 or custom_h1 > 0:
        return []
    chapter_pat = re.compile(r"^第[一二三四五六七八九十]+章\s+\S")
    section_pat = re.compile(r"^\d+\.\d+\s")  # "1.1 X" 节标题
    paras = manifest.get("paragraphs", [])

    # 找 body 中 "第N章" 段 (这些是有 prefix 的, anchor 自身)
    candidates: List[Dict] = []
    for i, p in enumerate(paras):
        if p.get("zone_guess") in ("toc", "cover"):
            continue
        text = p.get("text", "").strip()
        if chapter_pat.match(text) and p.get("style_name") not in {"Heading 1", "1-1级"}:
            # 找 anchor (默认 = 第N章段自身), 但若前面有 chapter 引言段, anchor 上移
            # 章引言段判定: 在 anchor 前 5 段窗口内, 段长 > 30, 不是 N.M section, 不是空, 是 body
            # 倒序扫 anchor 之前最多 6 段, 找紧贴 chapter heading 的章引言段 (连续长 body 段)
            # 撞上空段 (chapter 边界) / 上一章 / 上一节 / TOC / cover → 停
            preface_idx = i
            for k in range(i - 1, max(-1, i - 6), -1):
                pk = paras[k]
                tk = pk.get("text", "").strip()
                if not tk:
                    break  # 空段 = chapter 边界, 停; preface_idx 已是最早引言段
                if section_pat.match(tk) or chapter_pat.match(tk):
                    break
                if pk.get("zone_guess") in ("toc", "cover"):
                    break
                if len(tk) >= 30 and pk.get("zone_guess") in ("body", "unknown"):
                    preface_idx = k  # 候选引言段, 继续上推 (允许多段连续引言)
                    continue
                # 短段 (< 30 字) 当作非引言, 不更新 preface_idx 但继续上推
            anchor = paras[preface_idx]
            candidates.append({
                "anchor_paragraph_id": anchor["id"],
                "anchor_text_match": anchor.get("text", "").strip(),
                "title": text,
                "self_paragraph_id": p["id"],  # "第N章" 段 itself, 留作未来引用
            })
            if len(candidates) >= 10:
                break
    if not candidates:
        return []

    ops: List[Dict] = []
    for c in candidates:
        ops.append({
            "id": "",
            "type": "inject_heading_before",
            "enabled": False,  # 默认 disabled, 需人工 review (新增可见文字, 谨慎)
            "risk_class": "B",
            "visible_text_change": True,  # W4-C: 修正, inject 新增可见文字
            "params": {
                "anchor_paragraph_id": c["anchor_paragraph_id"],
                "anchor_text_match": c["anchor_text_match"],
                "title": c["title"],
                "level": 1,
            },
            "evidence": [
                f"body has '第N章' text but no Heading 1 / custom heading style",
                f"anchor = preface paragraph (W4-C)" if c["anchor_paragraph_id"] != c["self_paragraph_id"]
                else "anchor = chapter heading itself (no preface found)",
            ],
            "expected_delta": {"paragraphs_added": 1},
            "confidence": 0.70,
            "target_paragraph_ids": [c["anchor_paragraph_id"]],
        })
    return ops


def _detector_orphan_title_paragraph(manifest: Dict) -> List[Dict]:
    """W4-C: 检测客户手写的"裸章名段" (与 inject_heading_before 即将插入的 chapter title 重复).

    场景 (case19 round 2b): 客户 docx 在某 N.1 段之前手写了"X章名"段 (无"第N章"前缀),
    例如 "空间管理政策工具组合优化路径". inject_heading_before 会插 "第六章 空间管理政策工具组合优化路径",
    与该段重复. 此 detector 输出 delete op 与 inject op 配对, apply 时同时清除.

    判定: paragraphs[].text 是某 inject_heading_before op 的 title 去掉"第N章 "前缀后的字面.
    用 inject ops 推 — 但这里独立运行, 重新跑 chapter pattern.
    """
    chapter_pat = re.compile(r"^(第[一二三四五六七八九十]+章)\s+(\S.*)$")
    naked_chapter_names = set()
    paras = manifest.get("paragraphs", [])
    # 第一遍: 收集 body "第N章 X" 段中的 X 部分 (即裸章名)
    for p in paras:
        if p.get("zone_guess") in ("toc", "cover"):
            continue
        text = p.get("text", "").strip()
        m = chapter_pat.match(text)
        if m:
            naked_chapter_names.add(m.group(2).strip())
    if not naked_chapter_names:
        return []
    # 第二遍: 找 body 段 text 全等于裸章名 (即客户手写的"X章名"段)
    ops: List[Dict] = []
    for p in paras:
        if p.get("zone_guess") in ("toc", "cover"):
            continue
        text = p.get("text", "").strip()
        if text in naked_chapter_names:
            # 跳过 "第N章 X" 段自身 (它的 text 含前缀, 不会等于纯章名)
            ops.append({
                "id": "",
                "type": "delete_orphan_title_paragraph",
                "enabled": False,
                "risk_class": "B",
                "visible_text_change": True,  # 删除可见文字
                "params": {
                    "target_paragraph_id": p["id"],
                    "target_text_match": text,
                },
                "evidence": [
                    f"paragraph text {text!r} matches a chapter name (no '第N章' prefix); "
                    f"likely customer's manual heading duplicate of inject_heading_before"
                ],
                "expected_delta": {"paragraphs_deleted": 1},
                "confidence": 0.85,
                "target_paragraph_ids": [p["id"]],
            })
            if len(ops) >= 10:
                break
    return ops


DETECTORS = {
    "register_heading_style": _detector_register_heading_style,
    "relabel_pstyle": _detector_relabel_pstyle,
    "inject_heading_before": _detector_inject_heading_before,
    "orphan_title_paragraph": _detector_orphan_title_paragraph,
}


# ============================================================
# Plan
# ============================================================

def cmd_plan(docx_path: str, output_path: str, detect_filter: Optional[List[str]] = None,
             pre_manifest_path: Optional[str] = None) -> Dict:
    """跑 detectors, 写 surgery_plan.json + pre_surgery_manifest.json."""
    manifest = sm.build_probe_manifest(docx_path)
    out_dir = os.path.dirname(os.path.abspath(output_path)) or "."
    os.makedirs(out_dir, exist_ok=True)

    if pre_manifest_path is None:
        pre_manifest_path = os.path.join(out_dir, "pre_surgery_manifest.json")
    sm.write_manifest(manifest, pre_manifest_path)

    detector_names = list(DETECTORS.keys())
    if detect_filter:
        detector_names = [n for n in detector_names if n in detect_filter]

    ops: List[Dict] = []
    op_idx = 1
    for name in detector_names:
        for op in DETECTORS[name](manifest):
            op["id"] = f"op{op_idx:03d}"
            ops.append(op)
            op_idx += 1

    plan = {
        "schema_version": PLAN_SCHEMA_VERSION,
        "created_at": _now_iso(),
        "source_docx_sha256": _sha256_file(docx_path),
        "pre_manifest_path": os.path.basename(pre_manifest_path),
        "pre_manifest_sha256": _sha256_file(pre_manifest_path),
        "policy": {
            "allow_visible_text_change": False,
            "allowed_risk_classes": ["B"],
        },
        "operations": ops,
    }

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(plan, f, ensure_ascii=False, indent=2)

    enabled_count = sum(1 for op in ops if op.get("enabled"))
    print(f"  📋 plan: {len(ops)} ops ({enabled_count} enabled) -> {output_path}")
    for op in ops:
        flag = "✓" if op.get("enabled") else "·"
        print(f"     {flag} {op['id']} {op['type']} (conf={op.get('confidence', 0):.2f})")
    return plan


# ============================================================
# Apply
# ============================================================

def _apply_register_heading_style(temp_path: str, params: Dict) -> Dict:
    """注册 Heading N 内建样式. 幂等 (已存在不抛). 返回 actual_delta dict."""
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    level = params["level"]
    name = f"Heading {level}"
    d = Document(temp_path)
    added = False
    try:
        d.styles[name]
    except KeyError:
        try:
            d.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH, builtin=True)
            added = True
        except ValueError:
            # already exists internally (latent style 等), 幂等忽略
            pass
    d.save(temp_path)
    return {"styles_added": [name] if added else []}


def _apply_relabel_pstyle(temp_path: str, params: Dict, target_para_ids: List[str]) -> Dict:
    """in-place pStyle relabel. 不动文字.

    target_para_ids 仅 advisory (来自 manifest, idx 跟 python-docx paragraphs 可能错位
    因 textbox 内嵌套 <w:p>). 实际改用 style.name 匹配, 凡命中 style_map key 全改.
    """
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    style_map = params["style_map"]
    d = Document(temp_path)
    counts: Dict[str, int] = {tgt: 0 for tgt in set(style_map.values())}
    for p in d.paragraphs:
        sn = p.style.name if p.style else ""
        if sn not in style_map:
            continue
        target = style_map[sn]
        try:
            new_style_id = d.styles[target].style_id
        except KeyError:
            continue
        pPr = p._p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p._p.insert(0, pPr)
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is None:
            pStyle = OxmlElement("w:pStyle")
            pPr.insert(0, pStyle)
        pStyle.set(qn("w:val"), new_style_id)
        counts[target] = counts.get(target, 0) + 1
    d.save(temp_path)
    return {"heading_counts": {k: v for k, v in counts.items() if v}}


def _apply_inject_heading_before(temp_path: str, params: Dict) -> Dict:
    """W4-C: 在 anchor 段之前插入 Heading N 标题段.

    params 必含:
      - anchor_paragraph_id: manifest 段 id (e.g. "p000130")
      - anchor_text_match: anchor 段当前 text (apply 时用 text 重新查找, 避 idx 偏移问题)
      - title: 待插入的 chapter title (e.g. "第六章 空间管理政策工具组合优化路径")
      - level: int (1/2/3, default 1)

    实现:
      1. 用 anchor_text_match 在当前 docx 找 anchor 段 (容多 inject 顺序偏移)
      2. 创建新 <w:p>, pPr 设 pStyle=Heading N, runs 含一个 <w:t>{title}
      3. anchor_p._p.addprevious(new_p)

    注意: visible_text_change=True (新增可见文字), apply 后 hash 不与 pre 等 — verify 跳过 hash 比较.
    """
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    title = params.get("title")
    level = int(params.get("level", 1))
    anchor_text_match = params.get("anchor_text_match", "").strip()
    if not title or not anchor_text_match:
        raise RuntimeError(f"inject_heading_before missing title / anchor_text_match: {params}")

    d = Document(temp_path)
    anchor_p = None
    for p in d.paragraphs:
        if p.text.strip() == anchor_text_match:
            anchor_p = p
            break
    if anchor_p is None:
        raise RuntimeError(f"anchor not found by text: {anchor_text_match!r}")

    # Resolve Heading N style id (caller should have register_heading_style first)
    style_name = f"Heading {level}"
    try:
        style_id = d.styles[style_name].style_id
    except KeyError:
        # 容错: 没有 Heading N 样式 → 用 Normal (后续 relabel 处理)
        style_id = None

    new_p = OxmlElement("w:p")
    if style_id:
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), style_id)
        pPr.append(pStyle)
        new_p.append(pPr)
    new_run = OxmlElement("w:r")
    new_t = OxmlElement("w:t")
    new_t.text = title
    new_t.set(qn("xml:space"), "preserve")
    new_run.append(new_t)
    new_p.append(new_run)

    anchor_p._p.addprevious(new_p)
    d.save(temp_path)
    return {"paragraphs_added": 1, "title": title, "level": level}


def _apply_delete_orphan_title_paragraph(temp_path: str, params: Dict) -> Dict:
    """W4-C: 删指定段 (用 text 匹配, 避 idx 偏移).

    params 必含:
      - target_text_match: 待删段当前 text (整段 strip 后等值匹配)
      - target_paragraph_id: manifest 段 id (advisory, 不严格)

    用法: 与 inject_heading_before 配对, 删客户原稿"裸章名段"
    (e.g. "空间管理政策工具组合优化路径" 与 inject 的"第六章 空间管理政策工具组合优化路径" 重复).
    """
    from docx import Document

    target_text = params.get("target_text_match", "").strip()
    if not target_text:
        raise RuntimeError(f"delete_orphan_title_paragraph missing target_text_match: {params}")

    d = Document(temp_path)
    deleted_count = 0
    for p in list(d.paragraphs):
        if p.text.strip() == target_text:
            p._p.getparent().remove(p._p)
            deleted_count += 1
            break  # 只删第一个匹配段, 避免误伤
    if deleted_count == 0:
        raise RuntimeError(f"orphan target not found by text: {target_text!r}")
    d.save(temp_path)
    return {"paragraphs_deleted": deleted_count, "text": target_text}


_APPLIERS = {
    "register_heading_style": lambda tp, op: _apply_register_heading_style(tp, op["params"]),
    "relabel_pstyle": lambda tp, op: _apply_relabel_pstyle(tp, op["params"], op["target_paragraph_ids"]),
    "inject_heading_before": lambda tp, op: _apply_inject_heading_before(tp, op["params"]),
    "delete_orphan_title_paragraph": lambda tp, op: _apply_delete_orphan_title_paragraph(tp, op["params"]),
}


def cmd_apply(docx_path: str, plan_path: str, backup_tag: str = "PRE_SURGERY",
              output_dir: Optional[str] = None) -> Dict:
    """事务原子性 apply."""
    with open(plan_path, encoding="utf-8") as f:
        plan = json.load(f)

    actual_sha = _sha256_file(docx_path)
    if actual_sha != plan["source_docx_sha256"]:
        raise RuntimeError(
            f"source docx sha256 mismatch: {actual_sha[:16]} vs plan {plan['source_docx_sha256'][:16]}"
        )

    if output_dir is None:
        output_dir = os.path.dirname(os.path.abspath(plan_path)) or "."

    today = datetime.now().strftime("%Y-%m-%d")
    src_dir = os.path.dirname(os.path.abspath(docx_path))
    src_name = os.path.basename(docx_path)
    backup_path = os.path.join(src_dir, src_name.replace(".docx", f".{backup_tag}_BACKUP_{today}.docx"))
    if not os.path.exists(backup_path):
        shutil.copy(docx_path, backup_path)

    pre_visible_hash = _visible_text_hash(docx_path)

    pid = os.getpid()
    temp_path = os.path.join(src_dir, src_name.replace(".docx", f".surgery_tmp_{pid}.docx"))
    shutil.copy(docx_path, temp_path)

    report = {
        "started_at": _now_iso(),
        "plan_path": os.path.basename(plan_path),
        "source_docx_sha256": actual_sha,
        "backup_path": os.path.basename(backup_path),
        "pre_visible_text_sha256": pre_visible_hash,
        "operations": [],
        "status": "running",
        "failed_op_id": None,
        "error": None,
    }

    try:
        for op in plan["operations"]:
            if not op.get("enabled"):
                report["operations"].append({
                    "op_id": op["id"], "type": op["type"], "status": "skipped",
                    "actual_delta": {},
                })
                continue
            applier = _APPLIERS.get(op["type"])
            if applier is None:
                raise RuntimeError(f"unknown operation type: {op['type']}")
            actual_delta = applier(temp_path, op)
            report["operations"].append({
                "op_id": op["id"], "type": op["type"], "status": "applied",
                "actual_delta": actual_delta,
            })
        post_visible_hash = _visible_text_hash(temp_path)
        report["post_visible_text_sha256"] = post_visible_hash
        any_visible_change_op = any(
            op.get("enabled") and op.get("visible_text_change")
            for op in plan["operations"]
        )
        if not any_visible_change_op and post_visible_hash != pre_visible_hash:
            raise RuntimeError(
                f"visible text hash changed unexpectedly: {pre_visible_hash[:16]} -> {post_visible_hash[:16]}"
            )

        # 替换 source
        shutil.copy(temp_path, docx_path)
        report["status"] = "succeeded"
        os.remove(temp_path)
    except Exception as e:
        report["status"] = "failed"
        report["error"] = str(e)
        # find failed op
        applied_ids = {r["op_id"] for r in report["operations"] if r["status"] == "applied"}
        for op in plan["operations"]:
            if op.get("enabled") and op["id"] not in applied_ids:
                report["failed_op_id"] = op["id"]
                break
        # rename temp
        if os.path.exists(temp_path):
            failed_path = temp_path.replace(
                f".surgery_tmp_{pid}.docx",
                f".surgery_FAILED_{pid}_{int(datetime.now().timestamp())}.docx"
            )
            os.rename(temp_path, failed_path)
            report["failed_temp_path"] = os.path.basename(failed_path)
    finally:
        report["finished_at"] = _now_iso()

    apply_report_path = os.path.join(output_dir, "surgery_apply_report.json")
    with open(apply_report_path, "w", encoding="utf-8") as f:
        json.dump(report, f, ensure_ascii=False, indent=2)

    if report["status"] == "succeeded":
        post_manifest = sm.build_probe_manifest(docx_path)
        sm.write_manifest(post_manifest, os.path.join(output_dir, "post_surgery_manifest.json"))
        print(f"  ✅ apply succeeded ({len([r for r in report['operations'] if r['status']=='applied'])} ops)")
    else:
        print(f"  ❌ apply failed: {report['error']}")

    return report


# ============================================================
# Verify
# ============================================================

def cmd_verify(docx_path: str, plan_path: str, output_dir: Optional[str] = None) -> Dict:
    """独立 verify, 基于 manifest delta + plan expected_delta. 不调 risk-router."""
    with open(plan_path, encoding="utf-8") as f:
        plan = json.load(f)
    if output_dir is None:
        output_dir = os.path.dirname(os.path.abspath(plan_path)) or "."

    pre_path = os.path.join(output_dir, plan["pre_manifest_path"])
    post_path = os.path.join(output_dir, "post_surgery_manifest.json")
    if not os.path.isfile(pre_path) or not os.path.isfile(post_path):
        raise RuntimeError(f"manifests missing: pre={pre_path} post={post_path}")

    with open(pre_path, encoding="utf-8") as f:
        pre = json.load(f)
    with open(post_path, encoding="utf-8") as f:
        post = json.load(f)

    report: Dict = {
        "started_at": _now_iso(),
        "plan_path": os.path.basename(plan_path),
        "operations": [],
        "overall_passed": True,
    }

    # Heading counts in post (case-insensitive: python-docx 用小写 "heading 1")
    post_heading_counts: Dict[str, int] = {}
    for p in post.get("paragraphs", []):
        sn = p.get("style_name", "").lower()
        if sn.startswith("heading "):
            key = "Heading " + sn.split()[-1]
            post_heading_counts[key] = post_heading_counts.get(key, 0) + 1

    for op in plan["operations"]:
        op_report: Dict = {
            "op_id": op["id"], "type": op["type"],
            "enabled": op.get("enabled"),
            "expected_delta": op.get("expected_delta", {}),
            "checks": [],
            "passed": True,
        }
        if not op.get("enabled"):
            op_report["passed"] = True  # 跳过的 op 不算 fail
            report["operations"].append(op_report)
            continue
        if op["type"] == "relabel_pstyle":
            expected = op["expected_delta"].get("heading_counts", {})
            for tgt, n in expected.items():
                actual = post_heading_counts.get(tgt, 0)
                ok = actual >= n  # 至少这么多 (post 可能含其他来源)
                op_report["checks"].append({
                    "name": f"heading_counts.{tgt}>={n}",
                    "actual": actual, "expected": n, "ok": ok,
                })
                if not ok:
                    op_report["passed"] = False
        elif op["type"] == "register_heading_style":
            level = op["params"]["level"]
            name = f"Heading {level}"
            # 看 post manifest 有 paragraphs 引用 name (case-insensitive)
            ok = any(p.get("style_name", "").lower() == name.lower()
                     for p in post.get("paragraphs", []))
            op_report["checks"].append({
                "name": f"style_used.{name}", "actual": ok, "expected": True, "ok": ok,
            })
            if not ok:
                op_report["passed"] = False
        if not op_report["passed"]:
            report["overall_passed"] = False
        report["operations"].append(op_report)

    report["finished_at"] = _now_iso()
    verify_path = os.path.join(output_dir, "surgery_verify_report.json")
    with open(verify_path, "w", encoding="utf-8") as f:
        json.dump(report, f, ensure_ascii=False, indent=2)

    if report["overall_passed"]:
        print(f"  ✅ verify passed ({len(report['operations'])} ops checked)")
    else:
        failed = [op["op_id"] for op in report["operations"] if not op["passed"]]
        print(f"  ❌ verify failed for ops: {failed}")
    return report


# ============================================================
# CLI
# ============================================================

def main():
    ap = argparse.ArgumentParser(description="docx_surgery — B 类结构修复事务引擎")
    sub = ap.add_subparsers(dest="cmd", required=True)

    ap_plan = sub.add_parser("plan")
    ap_plan.add_argument("docx")
    ap_plan.add_argument("--output", required=True)
    ap_plan.add_argument("--detect", default="", help="comma-separated detector list")

    ap_apply = sub.add_parser("apply")
    ap_apply.add_argument("docx")
    ap_apply.add_argument("--plan", required=True)
    ap_apply.add_argument("--backup-tag", default="PRE_SURGERY")
    ap_apply.add_argument("--output-dir", default="")

    ap_verify = sub.add_parser("verify")
    ap_verify.add_argument("docx")
    ap_verify.add_argument("--plan", required=True)
    ap_verify.add_argument("--output-dir", default="")

    args = ap.parse_args()

    if args.cmd == "plan":
        detect = [d.strip() for d in args.detect.split(",") if d.strip()]
        cmd_plan(args.docx, args.output, detect_filter=detect or None)
    elif args.cmd == "apply":
        out = args.output_dir or None
        cmd_apply(args.docx, args.plan, backup_tag=args.backup_tag, output_dir=out)
    elif args.cmd == "verify":
        out = args.output_dir or None
        cmd_verify(args.docx, args.plan, output_dir=out)


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8")
    main()
