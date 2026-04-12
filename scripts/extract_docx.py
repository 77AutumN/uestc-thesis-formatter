#!/usr/bin/env python3
"""
extract_docx.py — 从 Word .docx 提取结构化内容并生成 LaTeX 章节文件

基于 UESTC 实战验证的提取引擎，吸收了两份 Gemini Review 的全部修复建议：
- NFKC 全角字符规范化
- 7 种标题正则模式容错
- 三层降级识别（样式 → 正则 → 关键词）
- 引用标记检测与报告

Usage:
    python extract_docx.py --input thesis.docx --output-dir ./output/
"""

import argparse
import json
import os
import re
import sys
import unicodedata

# === 正则模式（吸收 Review 的 7 条修复) ===

def normalize_text(text):
    """全角字符规范化 (Review #7: NFKC normalize)"""
    return unicodedata.normalize('NFKC', text)

# 一级标题: "第X章" (中文数字)
RE_CHAPTER_CN = re.compile(r'^\s*第[一二三四五六七八九十百]+章\s+(.*)', re.UNICODE)

# 二级标题: "X.Y 标题" (Review #1,#2,#3: 容忍无空格、前导缩进)
RE_SECTION = re.compile(r'^\s*(\d+)\.(\d+)\s*(.*)')

# 三级标题: "X.Y.Z 标题"
RE_SUBSECTION = re.compile(r'^\s*(\d+)\.(\d+)\.(\d+)\s*(.*)')

# 四级标题: "(X) 标题" 全角/半角括号 (Review #5)
RE_PARAGRAPH = re.compile(r'^\s*[（(](\d+)[）)]\s*(.*)')

# 中文数字序号: "一、问题的提出" (Review #7b)
RE_CN_NUMBERED = re.compile(r'^\s*[一二三四五六七八九十]+、\s*(.*)')


def parse_line(raw_line):
    """解析一行原始文本，返回 (type, content)"""
    raw_line = raw_line.rstrip('\r\n')
    if raw_line.startswith('[EMPTY]'):
        return ('empty', '')
    elif raw_line.startswith('[H1] '):
        return ('h1', raw_line[5:].strip())
    elif raw_line.startswith('[P] '):
        return ('paragraph', raw_line[4:].strip())
    else:
        return ('unknown', raw_line.strip())


def classify_paragraph(text):
    """
    对段落文本进行标题层级分类
    返回 (level, title_text) 或 None
    """
    normalized = normalize_text(text)

    # 章标题
    m = RE_CHAPTER_CN.match(normalized)
    if m:
        return ('chapter', m.group(1).strip())

    # 三级标题 (必须在二级之前检查)
    m = RE_SUBSECTION.match(normalized)
    if m:
        return ('subsection', m.group(4).strip() if m.group(4).strip() else normalized.strip())

    # 二级标题
    m = RE_SECTION.match(normalized)
    if m:
        return ('section', m.group(3).strip() if m.group(3).strip() else normalized.strip())

    # 四级标题 （X） — 已禁用：参考 PDF 最深只到 3 级 (X.X.X)
    # Word 中 （1）（2） 等为 Normal 样式，不应映射为 \subsubsection
    # m = RE_PARAGRAPH.match(normalized)
    # if m:
    #     return ('subsubsection', m.group(2).strip() if m.group(2).strip() else normalized.strip())

    return None


def escape_latex(text):
    """转义 LaTeX 特殊字符（保守策略）"""
    text = text.replace('&', '\\&')
    text = text.replace('%', '\\%')
    text = text.replace('#', '\\#')
    text = text.replace('_', '\\_')
    return text


def extract_from_docx(docx_path):
    """
    从 .docx 提取纯文本（带结构标记）
    Returns: list of (line_type, content)
    """
    try:
        from docx import Document
    except ImportError:
        print("❌ 缺少 python-docx 库，请运行: pip install python-docx")
        sys.exit(1)

    doc = Document(docx_path)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append(('empty', ''))
            continue

        # 检查段落样式
        style_name = para.style.name if para.style else ''
        if 'Heading 1' in style_name or '标题 1' in style_name:
            lines.append(('h1', text))
        elif 'Heading' in style_name or '标题' in style_name:
            lines.append(('heading', text))
        else:
            lines.append(('paragraph', text))

    return lines


def detect_citation_markers(lines):
    """检测正文中是否存在引用标记 [数字]"""
    count = 0
    for line_type, content in lines:
        if line_type in ('paragraph', 'unknown'):
            count += len(re.findall(r'\[\d+\]', content))
    return count


def find_chapter_boundaries(parsed_lines):
    """自动检测章节边界"""
    chapters = []
    for idx, (line_type, content) in enumerate(parsed_lines):
        if line_type == 'empty' or not content.strip():
            continue
        normalized = normalize_text(content)
        m = RE_CHAPTER_CN.match(normalized)
        if m or (line_type == 'h1' and re.search(r'第[一二三四五六七八九十百]+章', normalized)):
            chapter_title = m.group(1).strip() if m else normalized
            # 清理标题中的章号
            chapter_title = re.sub(r'^第[一二三四五六七八九十百]+章\s*', '', chapter_title)
            chapters.append({
                'idx': idx,
                'raw_title': content,
                'latex_title': chapter_title,
                'filename': f"ch{len(chapters)+1:02d}.tex"
            })
    return chapters


def find_special_sections(parsed_lines):
    """查找摘要、致谢、参考文献、攻读成果等特殊部分"""
    sections = {}
    for idx, (line_type, content) in enumerate(parsed_lines):
        normalized = normalize_text(content).strip()
        if '致' in normalized and '谢' in normalized and len(normalized) <= 10:
            sections['acknowledgement'] = idx
        elif '参 考 文 献' in normalized or '参考文献' in normalized:
            sections['references'] = idx
        elif '攻读' in normalized and '成果' in normalized:
            sections['accomplishments'] = idx
        elif normalized in ('摘 要', '摘要'):
            sections['abstract_zh'] = idx
        elif normalized == 'ABSTRACT':
            sections['abstract_en'] = idx
    return sections


def generate_chapter_tex(parsed_lines, start_idx, end_idx, latex_title):
    """生成单章的 .tex 内容"""
    tex_lines = [f"\\chapter{{{latex_title}}}\n"]

    for idx in range(start_idx + 1, end_idx):
        line_type, content = parsed_lines[idx]
        if line_type == 'empty' or not content.strip():
            continue

        heading = classify_paragraph(content)
        if heading:
            level, title = heading
            if level == 'chapter':
                continue  # 跳过重复章标题
            elif level == 'section':
                tex_lines.append(f"\n\\section{{{escape_latex(title)}}}\n")
            elif level == 'subsection':
                tex_lines.append(f"\n\\subsection{{{escape_latex(title)}}}\n")
            elif level == 'subsubsection':
                tex_lines.append(f"\n\\subsubsection{{{escape_latex(title)}}}\n")
        else:
            tex_lines.append(f"\n{escape_latex(content)}\n")

    return ''.join(tex_lines)


def extract_text_block(parsed_lines, start_idx, end_idx):
    """提取一段纯文本（用于摘要、致谢等）"""
    lines = []
    for idx in range(start_idx + 1, end_idx):
        _, content = parsed_lines[idx]
        if content.strip():
            lines.append(content.strip())
    return '\n\n'.join(lines)


def extract_cover_metadata(docx_path):
    """从 Word 封面表格中提取论文元数据（标题、作者、导师等）
    
    UESTC 论文模板封面使用表格布局：
      Table 0: 中文封面（论文题目、作者姓名、指导教师、学科专业、学院）
      Table 1: 第二封面（分类号、密级、完整标题、学位级别）
      Table 2: 英文封面（English title, Author, Supervisor, Discipline, School）
    
    采用标签匹配策略，从标签单元格的相邻单元格提取值。
    """
    try:
        from docx import Document
    except ImportError:
        print("❌ 缺少 python-docx 库")
        return {}

    doc = Document(docx_path)
    meta = {}

    # === 中文封面：从 Table 0 提取 ===
    if len(doc.tables) >= 1:
        t0 = doc.tables[0]
        for row in t0.rows:
            cells = [c.text.strip().replace('\u3000', ' ') for c in row.cells]
            # 去重（合并单元格会导致同一内容出现多次）
            unique_cells = []
            seen = set()
            for c in cells:
                if c and c not in seen:
                    unique_cells.append(c)
                    seen.add(c)

            for i, cell_text in enumerate(unique_cells):
                if '论文题目' in cell_text and i + 1 < len(unique_cells):
                    # 题目可能在同行的下一个单元格
                    title_part = unique_cells[i + 1]
                    meta['title_cn_part1'] = title_part
                elif '作者姓名' in cell_text and i + 1 < len(unique_cells):
                    meta['author_cn'] = unique_cells[i + 1].replace(' ', '')
                elif '指导教师' in cell_text and i + 1 < len(unique_cells):
                    raw = unique_cells[i + 1].replace(' ', '')
                    # 拆分姓名和职称："张三教授" 或 "张三　教授"
                    meta['advisor_cn_raw'] = unique_cells[i + 1]
                elif '学科专业' == cell_text or '学科专业' in cell_text:
                    if i + 1 < len(unique_cells) and '学科' not in unique_cells[i + 1]:
                        meta['major_cn'] = unique_cells[i + 1]
                elif cell_text == '学\u3000\u3000院' or cell_text == '学院' or cell_text.replace(' ', '') == '学院':
                    if i + 1 < len(unique_cells):
                        meta['school_cn'] = unique_cells[i + 1]

        # 题目可能跨两行，检查下一行是否是题目续行
        title_parts = []
        found_title_label = False
        for row in t0.rows:
            cells_unique = list(dict.fromkeys([c.text.strip() for c in row.cells if c.text.strip()]))
            if any('论文题目' in c for c in cells_unique):
                found_title_label = True
                for c in cells_unique:
                    if '论文题目' not in c and c:
                        title_parts.append(c)
            elif found_title_label and title_parts:
                # 紧跟题目行的下一行可能是续行
                row_texts = list(dict.fromkeys([c.text.strip() for c in row.cells if c.text.strip()]))
                if row_texts and not any(k in row_texts[0] for k in ['学科', '学号', '作者', '指导', '学院']):
                    title_parts.extend(row_texts)
                found_title_label = False  # 只取一行续行
            else:
                found_title_label = False
        if title_parts:
            meta['title_cn'] = ''.join(title_parts)

    # === 英文封面：从 Table 2 提取 ===
    if len(doc.tables) >= 3:
        t2 = doc.tables[2]
        # 英文封面通常是 2 列表格：[Label, Value]
        for row in t2.rows:
            cells = [c.text.strip() for c in row.cells]
            unique = list(dict.fromkeys([c for c in cells if c]))

            if len(unique) >= 1:
                first = unique[0]
                # 英文标题通常在第一行（无标签）
                if 'Master Thesis' in first or 'Submitted' in first:
                    continue  # 跳过 boilerplate
                if 'Author' == first and len(unique) >= 2:
                    meta['author_en'] = unique[1]
                elif 'Supervisor' == first and len(unique) >= 2:
                    meta['advisor_en'] = unique[1]
                elif 'Discipline' == first and len(unique) >= 2:
                    meta['major_en'] = unique[1]
                elif 'School' == first and len(unique) >= 2:
                    meta['school_en'] = unique[1]
                elif 'Student ID' == first and len(unique) >= 2:
                    meta['student_id'] = unique[1]

        # 英文标题：通常是 Table 2 的第一行
        first_row = t2.rows[0]
        first_cells = list(dict.fromkeys([c.text.strip() for c in first_row.cells if c.text.strip()]))
        if first_cells and 'Author' not in first_cells[0] and 'Discipline' not in first_cells[0]:
            meta['title_en'] = first_cells[0]

    # === 补充：从 Table 1 提取（第二封面有完整中文标题）===
    if len(doc.tables) >= 2 and 'title_cn' not in meta:
        t1 = doc.tables[1]
        for row in t1.rows:
            cells_unique = list(dict.fromkeys([c.text.strip() for c in row.cells if c.text.strip()]))
            if len(cells_unique) == 1 and len(cells_unique[0]) > 10:
                text = cells_unique[0]
                if '题名' not in text and '注' not in text and '学位' not in text:
                    meta['title_cn'] = text
                    break

    # === 清理导师信息 ===
    if 'advisor_cn_raw' in meta:
        raw = meta['advisor_cn_raw']
        # 常见格式: "张三　　教　授" 或 "张三 教授"
        import re
        cleaned = re.sub(r'[\s\u3000]+', ' ', raw).strip()
        parts = cleaned.split(' ')
        if len(parts) >= 2:
            meta['advisor_name_cn'] = parts[0]
            meta['advisor_title_cn'] = ' '.join(parts[1:])
        else:
            meta['advisor_name_cn'] = cleaned
            meta['advisor_title_cn'] = ''
        del meta['advisor_cn_raw']
    if 'title_cn_part1' in meta and 'title_cn' not in meta:
        meta['title_cn'] = meta['title_cn_part1']
    meta.pop('title_cn_part1', None)

    return meta


def main():
    parser = argparse.ArgumentParser(description='从 Word .docx 提取论文内容并生成 LaTeX 章节文件')
    parser.add_argument('--input', required=True, help='输入 .docx 文件路径')
    parser.add_argument('--output-dir', required=True, help='输出目录')
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"❌ 文件不存在: {args.input}")
        sys.exit(1)

    output_dir = args.output_dir
    chapters_dir = os.path.join(output_dir, 'chapters')
    os.makedirs(chapters_dir, exist_ok=True)

    # === Step 1: 提取 ===
    print(f"📖 正在提取: {args.input}")
    parsed_lines = extract_from_docx(args.input)
    print(f"  共 {len(parsed_lines)} 个段落")

    # === Step 2: 检测引用标记 ===
    citation_count = detect_citation_markers(parsed_lines)
    print(f"  正文引用标记 [数字]: {citation_count} 处")

    # === Step 3: 识别章节边界 ===
    chapters = find_chapter_boundaries(parsed_lines)
    print(f"  识别到 {len(chapters)} 个章节:")
    for ch in chapters:
        print(f"    📌 {ch['filename']}: {ch['raw_title']}")

    # === Step 4: 识别特殊部分 ===
    special = find_special_sections(parsed_lines)
    print(f"  特殊部分: {list(special.keys())}")

    # === Step 5: 输出 outline.json ===
    outline = {
        'chapters': [
            {'filename': ch['filename'], 'title': ch['raw_title'], 'latex_title': ch['latex_title']}
            for ch in chapters
        ],
        'special_sections': {k: True for k in special.keys()}
    }
    outline_path = os.path.join(output_dir, 'outline.json')
    with open(outline_path, 'w', encoding='utf-8') as f:
        json.dump(outline, f, ensure_ascii=False, indent=2)
    print(f"  ✅ outline.json")

    # === Step 6: 输出 thesis_meta.json ===
    meta = {
        'total_paragraphs': len([l for l in parsed_lines if l[0] != 'empty']),
        'total_chapters': len(chapters),
        'citation_markers_in_body': citation_count,
        'has_abstract_zh': 'abstract_zh' in special,
        'has_abstract_en': 'abstract_en' in special,
        'has_acknowledgement': 'acknowledgement' in special,
        'has_references': 'references' in special,
        'has_accomplishments': 'accomplishments' in special,
    }
    meta_path = os.path.join(output_dir, 'thesis_meta.json')
    with open(meta_path, 'w', encoding='utf-8') as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)
    print(f"  ✅ thesis_meta.json")

    # === Step 6.5: 提取封面元数据 ===
    print(f"\n  📋 提取封面元数据...")
    cover_meta = extract_cover_metadata(args.input)
    if cover_meta:
        cover_path = os.path.join(output_dir, 'cover_metadata.json')
        with open(cover_path, 'w', encoding='utf-8') as f:
            json.dump(cover_meta, f, ensure_ascii=False, indent=2)
        print(f"  ✅ cover_metadata.json")
        for k, v in cover_meta.items():
            print(f"    {k}: {v}")
    else:
        print(f"  ⚠️ 未能提取封面元数据，将使用默认值")

    # === Step 7: 生成章节 .tex 文件 ===
    for ch_idx, ch_info in enumerate(chapters):
        if ch_idx + 1 < len(chapters):
            end_idx = chapters[ch_idx + 1]['idx']
        elif 'acknowledgement' in special:
            end_idx = special['acknowledgement']
        elif 'references' in special:
            end_idx = special['references']
        else:
            end_idx = len(parsed_lines)

        tex_content = generate_chapter_tex(parsed_lines, ch_info['idx'], end_idx, ch_info['latex_title'])
        filepath = os.path.join(chapters_dir, ch_info['filename'])
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(tex_content)
        print(f"  ✅ {ch_info['filename']}")

    # === Step 8: 提取致谢 ===
    if 'acknowledgement' in special:
        ack_end = special.get('references', special.get('accomplishments', len(parsed_lines)))
        ack_text = extract_text_block(parsed_lines, special['acknowledgement'], ack_end)
        with open(os.path.join(output_dir, 'acknowledgement.txt'), 'w', encoding='utf-8') as f:
            f.write(ack_text)
        print(f"  ✅ acknowledgement.txt")

    # === Step 9: 提取参考文献原文 ===
    if 'references' in special:
        ref_end = special.get('accomplishments', len(parsed_lines))
        ref_text = extract_text_block(parsed_lines, special['references'], ref_end)
        with open(os.path.join(output_dir, 'references_raw.txt'), 'w', encoding='utf-8') as f:
            f.write(ref_text)
        print(f"  ✅ references_raw.txt")

    # === Step 10: 提取摘要 ===
    if 'abstract_zh' in special:
        abs_end = special.get('abstract_en', chapters[0]['idx'] if chapters else len(parsed_lines))
        abs_text = extract_text_block(parsed_lines, special['abstract_zh'], abs_end)
        with open(os.path.join(output_dir, 'abstract_zh.txt'), 'w', encoding='utf-8') as f:
            f.write(abs_text)
        print(f"  ✅ abstract_zh.txt")

    if 'abstract_en' in special:
        abs_end = chapters[0]['idx'] if chapters else len(parsed_lines)
        abs_text = extract_text_block(parsed_lines, special['abstract_en'], abs_end)
        with open(os.path.join(output_dir, 'abstract_en.txt'), 'w', encoding='utf-8') as f:
            f.write(abs_text)
        print(f"  ✅ abstract_en.txt")

    print(f"\n🎉 提取完成! 输出目录: {output_dir}")


if __name__ == '__main__':
    main()
