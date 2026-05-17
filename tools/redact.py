"""redact.py — PII redaction tool for OSS public release.

Implements the substitution rules from `docs/redaction-spec.md`:
- 13-digit student IDs        → <STUDENT_ID>
- Private real names          → CASE-A   (list loaded at runtime, see below)
- Absolute Windows paths      → ./
- Internal CASE-NNN codes     → CASE-A

Usage:
    python tools/redact.py --in-place .
    python tools/redact.py --check .       # exit 1 if any pattern still found

Configuration — private name list:
    Real student/advisor names are NOT hardcoded in this file. They live in
    `.redact-names.local.txt` at the repo root, which is gitignored and never
    committed. Format: one CJK or ASCII name per line; lines starting with `#`
    are ignored.

    On a clean OSS clone the file is absent and the name-replacement pass is a
    no-op (high-entropy patterns like student IDs / paths still run). In CI,
    the workflow re-creates the file from the `REDACT_NAMES` repo secret.

Scope:
    Text:   **/*.py, **/*.md, **/*.json, **/*.yaml, **/*.yml, **/*.txt, **/*.tex.
    Binary: **/*.docx (python-docx paragraph/table/core_properties text),
            **/*.pdf  (PyMuPDF page text). --check only — --in-place will not
            rewrite binary files; on hit it instructs the user to regen the
            source build script.
    Skips:  vendor/, .git/, .pytest_cache/, __pycache__/, tools/, .agent/,
            .codex_ipc/, .gemini_ipc/.
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

_PRIVATE_NAMES_FILE = Path(__file__).resolve().parent.parent / ".redact-names.local.txt"


def _load_private_names() -> list[str]:
    """Read the gitignored private-name list. Returns [] if the file is absent."""
    if not _PRIVATE_NAMES_FILE.is_file():
        return []
    names: list[str] = []
    for raw in _PRIVATE_NAMES_FILE.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if line and not line.startswith("#"):
            names.append(line)
    return names


def _build_patterns() -> list[tuple[re.Pattern[str], str]]:
    pats: list[tuple[re.Pattern[str], str]] = [
        (re.compile(r"\b20\d{11}\b"), "<STUDENT_ID>"),
    ]
    names = _load_private_names()
    cjk_names = [n for n in names if re.search(r"[一-鿿]", n)]
    ascii_names = [n for n in names if n not in cjk_names]
    if cjk_names:
        pats.append((re.compile("|".join(re.escape(n) for n in cjk_names)), "CASE-A"))
    if ascii_names:
        pats.append((
            re.compile(r"\b(?:" + "|".join(re.escape(n) for n in ascii_names) + r")\b",
                       re.IGNORECASE),
            "CASE-A",
        ))
    pats.extend([
        # Collapse CASE-NNN/NNN/NNN sequences first so partial-redacts don't leave dangling.
        (re.compile(r"\bCASE-0\d{2}(?:[/,\s]+0\d{2})+\b"), "CASE-A"),
        (re.compile(r"\bCASE-A[/,\s]+0\d{2}(?:[/,\s]+0\d{2})*\b"), "CASE-A"),
        # Drive-letter path (case-insensitive D:/d:; non-greedy class avoids eating delimiters).
        (re.compile(r"[Dd]:[\\/]+[Oo]pen claw[\\/][^\s'\"`,)]+"), "./"),
        # Bare project-name reference (no drive prefix) — collapse to ./ as well.
        (re.compile(r"\b[Oo]pen claw\b[\\/]?[^\s'\"`,)]*"), "./"),
        # Internal case codes: numeric CASE-NNN, alpha CASE-LETTERS (exclude CASE-A/B sentinels).
        (re.compile(r"\bCASE-(?!A\b|B\b)[A-Z]{2,}\b"), "CASE-A"),
        (re.compile(r"\bCASE-0\d{2}\b"), "CASE-A"),
        # Lowercase case-id forms (case011, _case015_round1_fixN_*) that bypass upper-case PATTERN.
        # `_` is a word-char so \b fails inside identifiers; use negative digit-boundaries.
        (re.compile(r"_case\d{3}_round\d+_[a-zA-Z0-9_]+"), "_case_anon"),
        (re.compile(r"(?<!\d)case0\d{2,3}(?!\d)"), "case_anon"),
    ])
    return pats


PATTERNS = _build_patterns()

TEXT_EXTS = {".py", ".md", ".json", ".yaml", ".yml", ".txt", ".tex"}
BINARY_EXTS = {".docx", ".pdf"}
ALL_EXTS = TEXT_EXTS | BINARY_EXTS
# Back-compat alias for any external caller.
EXTS = TEXT_EXTS
# Narrowed: only skip directories that genuinely contain pattern literals or are
# git/cache scaffolding. We DO want to scan .github/*.md (e.g. copilot-instructions.md).
SKIP_DIRS = {"vendor", ".git", ".pytest_cache", "__pycache__", "tools",
             ".agent", ".codex_ipc", ".gemini_ipc"}
# File-level skip: the workflow + redaction-spec themselves contain pattern literals.
SKIP_FILES = {
    ".github/workflows/redact-check.yml",
    "docs/redaction-spec.md",
    ".redact-names.local.txt",
}


def _is_skipped_file(p: Path, root: Path) -> bool:
    try:
        rel = p.relative_to(root).as_posix()
    except ValueError:
        return False
    return rel in SKIP_FILES


def iter_files(root: Path):
    for p in root.rglob("*"):
        if not p.is_file():
            continue
        if p.suffix.lower() not in ALL_EXTS:
            continue
        if any(part in SKIP_DIRS for part in p.parts):
            continue
        if _is_skipped_file(p, root):
            continue
        yield p


def _extract_docx_text(path: Path) -> str:
    """Extract visible text from a .docx (paragraphs + tables + core properties).

    Lazy import so the tool still runs on clones that don't have python-docx
    installed — they just won't be able to scan .docx fixtures, which the
    caller reports as a soft failure.
    """
    from docx import Document  # type: ignore[import-not-found]
    d = Document(str(path))
    parts: list[str] = [p.text for p in d.paragraphs]
    for tbl in d.tables:
        for row in tbl.rows:
            parts.extend(c.text for c in row.cells)
    cp = d.core_properties
    parts.extend([
        cp.author or "", cp.title or "", cp.subject or "",
        cp.keywords or "", cp.comments or "", cp.last_modified_by or "",
    ])
    return "\n".join(p for p in parts if p)


def _extract_pdf_text(path: Path) -> str:
    """Extract text from a .pdf via PyMuPDF page-by-page get_text."""
    import fitz  # type: ignore[import-not-found]
    doc = fitz.open(str(path))
    try:
        return "\n".join(page.get_text() for page in doc)
    finally:
        doc.close()


BINARY_EXTRACTORS = {".docx": _extract_docx_text, ".pdf": _extract_pdf_text}


def redact_text(text: str) -> tuple[str, int]:
    total = 0
    for pat, repl in PATTERNS:
        text, n = pat.subn(repl, text)
        total += n
    return text, total


def cmd_inplace(root: Path) -> int:
    grand_total = 0
    files_changed = 0
    binary_skipped = 0
    for f in iter_files(root):
        ext = f.suffix.lower()
        if ext in BINARY_EXTS:
            # Binary rewrite is unsafe (docx/pdf structure). Skip silently
            # during --in-place; --check is responsible for flagging hits.
            binary_skipped += 1
            continue
        try:
            original = f.read_text(encoding="utf-8")
        except (UnicodeDecodeError, OSError):
            continue
        new, n = redact_text(original)
        if n:
            f.write_text(new, encoding="utf-8")
            print(f"[redact] {f.relative_to(root)}: {n} replacement(s)")
            grand_total += n
            files_changed += 1
    print(f"\nDone: {grand_total} replacement(s) across {files_changed} file(s)"
          + (f" ({binary_skipped} binary file(s) skipped — use --check to scan, regen via build script)"
             if binary_skipped else ""))
    return 0


def cmd_check(root: Path) -> int:
    hits: list[tuple[Path, int, str]] = []
    binary_errors: list[tuple[Path, str]] = []
    for f in iter_files(root):
        ext = f.suffix.lower()
        if ext in BINARY_EXTS:
            extractor = BINARY_EXTRACTORS[ext]
            try:
                text = extractor(f)
            except ImportError as e:
                binary_errors.append((f.relative_to(root),
                                      f"missing dependency for {ext}: {e}"))
                continue
            except Exception as e:  # noqa: BLE001 — surface any extractor failure
                binary_errors.append((f.relative_to(root),
                                      f"extractor failed: {type(e).__name__}: {e}"))
                continue
        else:
            try:
                text = f.read_text(encoding="utf-8")
            except (UnicodeDecodeError, OSError):
                continue
        for line_no, line in enumerate(text.splitlines(), 1):
            for pat, _ in PATTERNS:
                if pat.search(line):
                    hits.append((f.relative_to(root), line_no, line.strip()[:120]))
                    break
    if binary_errors:
        # Soft warning, but for binary files where we can't extract text we
        # cannot guarantee CLEAN — emit a warning so reviewers know to install
        # the binary dep before trusting the result.
        for path, msg in binary_errors:
            print(f"::warning::redact.py could not scan {path}: {msg}",
                  file=sys.stderr)
    if hits:
        print(f"BLOCKED — {len(hits)} PII hit(s):", file=sys.stderr)
        for path, line_no, snippet in hits:
            print(f"  {path}:{line_no}: {snippet}", file=sys.stderr)
        return 1
    if binary_errors:
        print(f"PARTIAL — no PII patterns found in scanned files, but "
              f"{len(binary_errors)} binary file(s) could not be scanned "
              f"(see warnings above). Install python-docx + PyMuPDF and rerun.",
              file=sys.stderr)
        return 2
    print("CLEAN — no PII patterns found")
    return 0


def cmd_audit_history(root: Path) -> int:
    """Audit git history with `git log -S/-G` for known PII patterns.

    Reports commits whose diffs introduce or remove a sensitive literal /
    regex match. Useful before publishing a release to verify that ALL
    history (not just HEAD) is clean. Exit 1 on any historical hit.

    Note: detects strings that EVER appeared in any commit on any branch —
    even if later removed. To purge dangling blobs after a hit, use
    `git filter-repo --replace-text` and force-push, then ask GitHub
    Support to garbage-collect (force-push alone won't unreach the blob).
    """
    import subprocess  # noqa: PLC0415 — lazy, only when audit-history runs

    # Build (search-mode, term) pairs. `-S` is fast literal pickaxe;
    # `-G` is regex (slower) — use only for patterns that have no clean
    # literal anchor.
    literal_terms = list(_load_private_names())
    literal_terms.extend([
        r"D:\open claw", r"d:\open claw",
        "Open claw", "open claw",
    ])
    regex_terms = [
        r"\b20[0-9]{11}\b",          # student ID
        r"\bCASE-[A-Z]{2,}\b",       # CASE-LETTERS codes
        r"\b_case[0-9]{3}_round",    # internal script naming
        r"(?<!\w)case0[0-9]{2,3}(?!\w)",  # lowercase case-id
    ]

    hits: list[tuple[str, str, str, str]] = []  # (mode, term, sha, subject)
    git_root = str(root.resolve())
    for term in literal_terms:
        if not term:
            continue
        cp = subprocess.run(
            ["git", "-C", git_root, "log", "--all", "-S", term,
             "--pretty=format:%H %s"],
            capture_output=True, text=False)
        out = cp.stdout.decode("utf-8", errors="replace").strip()
        if out:
            for line in out.splitlines():
                sha, _, subj = line.partition(" ")
                hits.append(("S", term, sha[:12], subj[:60]))
    for term in regex_terms:
        cp = subprocess.run(
            ["git", "-C", git_root, "log", "--all", "-G", term,
             "--pretty=format:%H %s"],
            capture_output=True, text=False)
        out = cp.stdout.decode("utf-8", errors="replace").strip()
        if out:
            for line in out.splitlines():
                sha, _, subj = line.partition(" ")
                hits.append(("G", term, sha[:12], subj[:60]))

    if hits:
        print(f"HISTORY AUDIT — {len(hits)} historical pattern hit(s):",
              file=sys.stderr)
        for mode, term, sha, subj in hits:
            print(f"  [-{mode} {term[:40]:40s}] {sha} {subj}", file=sys.stderr)
        print("\nNote: these are HISTORICAL hits. Current HEAD may be clean.",
              file=sys.stderr)
        print("To purge, see docs/redaction-spec.md (git filter-repo + GitHub "
              "Support cache invalidation).", file=sys.stderr)
        return 1

    print("HISTORY CLEAN — no historical pattern hits across all branches.")
    return 0


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("path", type=Path, help="Root directory to scan")
    grp = ap.add_mutually_exclusive_group(required=True)
    grp.add_argument("--in-place", action="store_true", help="Apply replacements to text files (binary skipped)")
    grp.add_argument("--check", action="store_true", help="Exit 1 if PII remains in working tree (incl. .docx/.pdf)")
    grp.add_argument("--audit-history", action="store_true",
                     help="Exit 1 if PII appears anywhere in git history (release-time gate)")
    args = ap.parse_args()
    if not args.path.is_dir():
        print(f"ERROR: {args.path} is not a directory", file=sys.stderr)
        return 2
    if args.in_place:
        return cmd_inplace(args.path)
    if args.audit_history:
        return cmd_audit_history(args.path)
    return cmd_check(args.path)


if __name__ == "__main__":
    sys.exit(main())
